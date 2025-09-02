from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from django.conf import settings
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
import os
import re
from .estatutos_sociedad import EstatutosSociedad
from num2words import num2words
from docx.oxml.ns import qn
from docx.oxml import OxmlElement 

def descargar_docx_estatutos_sociedad(request, pk):
    """
    Genera y descarga un archivo DOCX de Estatutos Sociales con placeholders reemplazados
    """
    # Obtener el objeto EstatutosSociedad
    estatutos = get_object_or_404(EstatutosSociedad, pk=pk, usuario=request.user)
    
    # Ruta al template DOCX - usando la misma estructura que Acta de Asamblea
    template_path = os.path.join(
        settings.BASE_DIR.parent, 
        "DOCUMENTOS OLEA ABOGADOS", 
        "Estatutos Sociales", 
        "Estatutos Sociales PLACE.docx"
    )
    
    if not os.path.exists(template_path):
        return HttpResponse("Template file not found", status=404)
    
    try:
        # Cargar el documento template
        doc = Document(template_path)
        
        # Diccionario de reemplazos - usando los placeholders exactos del DOCX
        replacements = {
            '{{denominacion}}': estatutos.denominacion or '',
            '{{DENOMINACION}}': estatutos.denominacion.upper() or '',
            #'{{DENOMINACION_BOLD}}': f"<b>{estatutos.denominacion.upper()}</b>" if estatutos.denominacion.upper() else '',
            '{{forma_legal}}': estatutos.forma_legal or '',
            '{{domicilio}}': estatutos.domicilio or '',
            '{{nacionalidad}}': estatutos.nacionalidad or '',
            '{{objeto_social}}': estatutos.objeto_social or '',
            '{{capital_fijo_monto}}': f"${estatutos.capital_fijo_monto:,.2f}" if estatutos.capital_fijo_monto else '',
            '{{capital_fijo_texto}}': estatutos.capital_fijo_texto or '',
            '{{acciones_serie_a}}': f"{estatutos.acciones_serie_a:,.0f}" if estatutos.acciones_serie_a else '',
            '{{acciones_serie_a_texto}}': num2words(estatutos.acciones_serie_a, lang="es").capitalize() if estatutos.acciones_serie_a else "",
        }
        
        # Función auxiliar para reemplazar texto completo en párrafos (copiada de Acta de Asamblea)
        def replace_in_paragraph(paragraph, replacements):
            # Obtener todo el texto del párrafo
            full_text = paragraph.text
            
            # Verificar si hay placeholders en el párrafo
            has_placeholder = False
            found_placeholder = None
            for placeholder in replacements.keys():
                if placeholder in full_text:
                    has_placeholder = True
                    found_placeholder = placeholder
                    break
            
            if not has_placeholder:
                return
            
            # Manejo especial para objeto_social con múltiples líneas y formato legal
            if found_placeholder == '{{objeto_social}}' and replacements[found_placeholder]:
                objeto_social_text = replacements[found_placeholder]
                
                # Verificar si es una lista o string y convertir a string si es necesario
                if isinstance(objeto_social_text, list):
                    objeto_social_text = '\n'.join(str(item) for item in objeto_social_text)
                elif not isinstance(objeto_social_text, str):
                    objeto_social_text = str(objeto_social_text)
                
                # Limpiar todos los runs existentes del párrafo original
                for run in paragraph.runs:
                    run.text = ""
                
                # Eliminar el placeholder del texto original
                new_text = full_text.replace(found_placeholder, "")
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                else:
                    paragraph.add_run(new_text)
                
                # Procesar el objeto social línea por línea
                lines = objeto_social_text.split('\n')
                parent = paragraph._element.getparent()
                current_p = paragraph._element
                
                # Contador para numeración automática
                item_number = 1
                
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    
                    # Crear nuevo párrafo para cada elemento
                    new_p = doc.add_paragraph()
                    
                    # Aplicar formato legal específico
                    pf = new_p.paragraph_format
                    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justificado
                    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    pf.line_spacing = 1.15  # Interlineado 1.15
                    pf.space_after = Pt(6)  # 6 pt después del párrafo
                    """
                    pf.left_indent = Inches(0.75)  # Sangría izquierda de 0.75 cm (aprox)
                    pf.first_line_indent = Inches(-0.75)  # Sangría francesa (negativa)
                    
                    # Agregar numeración y texto
                    run = new_p.add_run(f"\t{line}")
                    
                    # Aplicar formato de fuente
                    font = run.font
                    font.name = 'Times New Roman'
                    font.size = Pt(12)
                    """
                    # --- SANGRÍA FRANCESA CON TAB STOP ---
                    # Texto alineado a 0.75 cm; número queda “antes” (en el margen)
                    pf.space_after = Pt(12)   # 12 pt = about one extra line
                    pf.space_before = Pt(0)
                    
                    texto_inicio = Cm(0.75)
                    pf.left_indent = Cm(2)
                    pf.first_line_indent = -Cm(1)

                    # Tab stop EXACTO en 0.75 cm para que el texto arranque alineado
                    pPr = new_p._p.get_or_add_pPr()
                    tabs = pPr.find(qn('w:tabs'))
                    if tabs is None:
                        tabs = OxmlElement('w:tabs')
                        pPr.append(tabs)
                    tab = OxmlElement('w:tab')
                    tab.set(qn('w:val'), 'left')
                    # w:pos espera twips (1 pt = 20 twips)
                    tab.set(qn('w:pos'), str(int(texto_inicio.pt * 20)))
                    tabs.append(tab)

                    # --- CONTENIDO ---
                    # El usuario ya mete el número. Forzamos una TAB tras el primer punto para separar número y texto.
                    txt = line
                    if '\t' not in txt:
                        if '. ' in txt:
                            txt = txt.replace('. ', '.\t', 1)
                        elif '.' in txt:
                            txt = txt.replace('.', '.\t', 1)

                    run = new_p.add_run(txt)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)

                    # Mover el párrafo a la posición correcta
                    new_p_element = new_p._element
                    parent.insert(parent.index(current_p) + item_number, new_p_element)
                    
                    item_number += 1
                
                return
            
            # Realizar reemplazos normales para otros placeholders
            new_text = full_text
            for placeholder, value in replacements.items():
                if placeholder in new_text:
                    new_text = new_text.replace(placeholder, str(value))
            
            # Si el texto cambió, reemplazar todo el párrafo
            if new_text != full_text:
                # Limpiar todos los runs existentes
                for run in paragraph.runs:
                    run.text = ""
                
                # Agregar el nuevo texto al primer run
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                else:
                    paragraph.add_run(new_text)
        
        # Reemplazar placeholders en párrafos
        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph, replacements)
        
        # Reemplazar placeholders en tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph, replacements)
        
        # Reemplazar placeholders en headers y footers
        for section in doc.sections:
            # Header
            if section.header:
                for paragraph in section.header.paragraphs:
                    replace_in_paragraph(paragraph, replacements)
            
            # Footer
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    replace_in_paragraph(paragraph, replacements)
        
        # Preparar la respuesta HTTP
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        # Nombre del archivo de descarga
        filename = f"Estatutos_Sociales_{estatutos.denominacion.replace(' ', '_') if estatutos.denominacion else 'documento'}.docx"
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        # Guardar el documento en la respuesta
        doc.save(response)
        
        return response
        
    except Exception as e:
        return HttpResponse(f"Error generating document: {str(e)}", status=500)
