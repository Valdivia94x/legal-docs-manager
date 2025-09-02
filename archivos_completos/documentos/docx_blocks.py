"""
Utility module for DOCX block manipulation and formatting.
Handles the insertion of complex formatted sections like "Orden del Día + Resoluciones".
"""

import json
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Pt, Cm


def to_roman(num):
    """Convert an integer to Roman numerals (I, II, III, IV, V, etc.)"""
    if num <= 0:
        return ""
    
    val = [1000, 900, 500, 400, 100, 90, 50, 40, 10, 9, 5, 4, 1]
    syms = ["M", "CM", "D", "CD", "C", "XC", "L", "XL", "X", "IX", "V", "IV", "I"]
    roman = ''
    i = 0
    while num > 0:
        for _ in range(num // val[i]):
            roman += syms[i]
            num -= val[i]
        i += 1
    return roman


def build_ordenes_con_resoluciones(orden_dia_json, resoluciones_json=None):
    """
    Build a unified list of agenda items with their resolutions.
    Ensures compatibility between nested format and separate resoluciones_json.
    
    Args:
        orden_dia_json: List of agenda items (may include nested resoluciones)
        resoluciones_json: Optional separate list of resolutions to merge by punto
    
    Returns:
        List of agenda items with resoluciones properly nested
    """
    if not orden_dia_json:
        return []
    
    # Parse JSON if it's a string
    if isinstance(orden_dia_json, str):
        try:
            orden_dia_json = json.loads(orden_dia_json)
        except (json.JSONDecodeError, TypeError):
            return []
    
    if isinstance(resoluciones_json, str):
        try:
            resoluciones_json = json.loads(resoluciones_json)
        except (json.JSONDecodeError, TypeError):
            resoluciones_json = None
    
    # Ensure orden_dia_json is a list
    if not isinstance(orden_dia_json, list):
        return []
    
    # Build the unified structure
    ordenes = []
    
    for item in orden_dia_json:
        if isinstance(item, dict):
            orden = {
                'numero': item.get('numero', len(ordenes) + 1),
                'titulo': item.get('titulo', ''),
                'descripcion': item.get('descripcion', ''),
                'resoluciones': item.get('resoluciones', [])
            }
        elif isinstance(item, str):
            orden = {
                'numero': len(ordenes) + 1,
                'titulo': item,
                'descripcion': '',
                'resoluciones': []
            }
        else:
            continue
        
        # If resoluciones_json exists and this orden doesn't have resoluciones,
        # try to find matching resoluciones by punto number
        if resoluciones_json and not orden['resoluciones']:
            if isinstance(resoluciones_json, list):
                for resolucion in resoluciones_json:
                    if isinstance(resolucion, dict):
                        punto = resolucion.get('punto')
                        if punto and (punto == orden['numero'] or str(punto) == str(orden['numero'])):
                            orden['resoluciones'].append(resolucion)
        
        ordenes.append(orden)
    
    return ordenes


def find_paragraph_with_placeholder(doc, placeholder):
    """
    Find the paragraph containing the placeholder, even if it's inside a table.
    
    Args:
        doc: Document object
        placeholder: String to search for (e.g., "{{ORDENES_Y_RESOLUCIONES}}")
    
    Returns:
        Tuple (paragraph, parent_element) where parent_element could be doc or a table cell
    """
    # Search in main document body
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            return paragraph, doc
    
    # Search in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if placeholder in paragraph.text:
                        return paragraph, cell
    
    return None, None


def inject_ordenes_y_resoluciones(doc, placeholder, ordenes, texto_constante):
    """
    Inject the formatted "Orden del Día + Resoluciones" section into the document.
    
    Args:
        doc: Document object
        placeholder: Placeholder string to replace
        ordenes: List of agenda items with resoluciones
        texto_constante: Constant text to insert before the detailed sections
    
    Raises:
        ValueError: If placeholder is not found in the document
    """
    # Find the paragraph with the placeholder
    target_paragraph, parent_element = find_paragraph_with_placeholder(doc, placeholder)
    
    if not target_paragraph:
        raise ValueError(f"Placeholder {placeholder} no encontrado en la plantilla.")
    
    # Get the position of the target paragraph
    if parent_element == doc:
        # Paragraph is in main document body
        try:
            paragraphs_list = doc.paragraphs
            target_index = paragraphs_list.index(target_paragraph)
        except ValueError:
            # Fallback: find paragraph by iterating through elements
            target_index = 0
            for i, p in enumerate(paragraphs_list):
                if p._element == target_paragraph._element:
                    target_index = i
                    break
        
        # Clear the placeholder paragraph
        target_paragraph.clear()
        
        # Insert content after the cleared paragraph
        insert_position = target_index
        
    else:
        # Paragraph is in a table cell
        try:
            paragraphs_list = parent_element.paragraphs
            target_index = paragraphs_list.index(target_paragraph)
        except ValueError:
            # Fallback: find paragraph by iterating through elements
            target_index = 0
            for i, p in enumerate(paragraphs_list):
                if p._element == target_paragraph._element:
                    target_index = i
                    break
        
        # Clear the placeholder paragraph
        target_paragraph.clear()
        
        # Insert content after the cleared paragraph
        insert_position = target_index
    
    # Helper function to insert paragraph with proper styling
    def insert_paragraph_at_position(pos, text="", alignment=WD_ALIGN_PARAGRAPH.LEFT, bold=False, style_name=None):
        if parent_element == doc:
            # Insert in main document
            new_p = doc.add_paragraph()
            # Move to correct position
            doc._body._body.insert(pos + 1, new_p._element)
        else:
            # Insert in table cell
            new_p = parent_element.add_paragraph()
        
        # Apply style if specified
        if style_name:
            try:
                new_p.style = doc.styles[style_name]
            except KeyError:
                # Fallback to default style if not found
                pass
        
        if text:
            run = new_p.add_run(text)
            run.bold = bold
        
        new_p.alignment = alignment
        return new_p
    
    current_pos = insert_position
    
    if not ordenes:
        # If no ordenes, just insert the constant text
        if texto_constante:
            insert_paragraph_at_position(current_pos, texto_constante, style_name="Body Text")
        return

    # 2. Insert list of all agenda items with List Paragraph style
    for orden in ordenes:
        """
        roman_num = to_roman(orden['numero'])
        titulo = orden.get('titulo', f"Punto {orden['numero']}")
        agenda_text = f"{roman_num}. {titulo}"
        # Inserta el título de la orden
        p=insert_paragraph_at_position(current_pos, agenda_text, style_name="Body Text")
        p.paragraph_format.space_after = Pt(18)
        # si este es el último punto, darle 38pt para separar del texto constante
        if orden == ordenes[-1]:
            p.paragraph_format.space_after = Pt(28)
        current_pos += 1
        """
        roman_num = to_roman(orden['numero'])
        titulo = orden.get('titulo', f"Punto {orden['numero']}")

        p = insert_paragraph_at_position(current_pos, "", style_name="Body Text")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # Sangría: primera línea “sale” a la izquierda (colgante)
        p.paragraph_format.left_indent = Cm(1.30)         # sangría general
        p.paragraph_format.first_line_indent = Cm(-1.30)  # colgante
        p.paragraph_format.space_after = Pt(18)

        # Número + tab + título
        p.add_run(f"{roman_num}. ").bold = False
        p.add_run("\t")           # tabulador (cae en la sangría)
        p.add_run(titulo)

        if orden == ordenes[-1]:
            p.paragraph_format.space_after = Pt(28)

        current_pos += 1
    
    # 3. Insert constant text with Body Text style
    if texto_constante:
        p=insert_paragraph_at_position(current_pos, texto_constante, 
                                   alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, style_name="Body Text")
        p.paragraph_format.space_after = Pt(38)
        current_pos += 1
    
    # 4. Insert detailed sections for each agenda item
    for i, orden in enumerate(ordenes):
        # Add blank line before each section (except first)
        if i > 0:
            insert_paragraph_at_position(current_pos, "", style_name="Body Text")
            current_pos += 1
        
        # Title as Heading 2 with Roman numeral
        """
        roman_num = to_roman(orden['numero'])
        titulo = orden.get('titulo', f"Punto {orden['numero']}")
        title_text = f"{roman_num}. {titulo}"
        insert_paragraph_at_position(current_pos, title_text, 
                                   alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, style_name="Heading 2")
        """
        roman_num = to_roman(orden['numero'])
        titulo = orden.get('titulo', f"Punto {orden['numero']}")

        p = insert_paragraph_at_position(current_pos, "", style_name="Body Text")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # Sangría: primera línea “sale” a la izquierda (colgante)
        p.paragraph_format.left_indent = Cm(1.30)         # sangría general
        p.paragraph_format.first_line_indent = Cm(-1.30)  # colgante
        p.paragraph_format.space_after = Pt(28)

        # Número + tab + título
        p.add_run(f"{roman_num}. ").bold = True
        p.add_run("\t")           # tabulador (cae en la sangría)
        p.add_run(titulo).bold = True
        current_pos += 1
        
        # Description with Body Text style and justified alignment
        descripcion = orden.get('descripcion', '').strip()
        if descripcion:
            """
            p = insert_paragraph_at_position(current_pos, descripcion, 
                                       alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, style_name="Body Text")
            p.paragraph_format.space_after = Pt(25)
            current_pos += 1
            """
            # Divide en párrafos usando doble salto de línea o simple \n
            # si en la BD guardas con saltos simples.
            for bloque in descripcion.split("\n"):
                bloque = bloque.strip()
                if not bloque:
                    continue  # salta líneas vacías
                p = insert_paragraph_at_position(
                    current_pos,
                    bloque,
                    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    style_name="Body Text"
                )
                p.paragraph_format.space_after = Pt(25)  # separación entre cada párrafo
                current_pos += 1
        
        # Resolution header (Heading 1, centered)
        """resolution_header = "R E S O L U C I Ó N".bold()
        p = insert_paragraph_at_position(current_pos, resolution_header, 
                                   alignment=WD_ALIGN_PARAGRAPH.CENTER, style_name="Heading 1")
        p.paragraph_format.space_after = Pt(25)
        current_pos += 1"""
        
        # Create an empty paragraph first
        p = insert_paragraph_at_position(
            current_pos,
            "",  # start with empty text
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            #style_name="Heading 1"
        )

        # Add the run with text and set bold
        run = p.add_run("R E S O L U C I Ó N")
        run.bold = True

        # Adjust spacing
        p.paragraph_format.space_after = Pt(25)
        current_pos += 1

        
        # Resolutions with Body Text style and justify alignment
        resoluciones = orden.get('resoluciones', [])
        if resoluciones:
            for resolucion in resoluciones:
                """
                if isinstance(resolucion, dict):
                    clave = resolucion.get('clave', '')
                    texto = resolucion.get('texto', '')
                    
                    if clave and texto:
                        res_text = f"{clave}. {texto}"
                    elif texto:
                        res_text = texto
                    else:
                        continue
                    
                    insert_paragraph_at_position(current_pos, res_text, 
                                               alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, style_name="Body Text")
                    current_pos += 1
                elif isinstance(resolucion, str) and resolucion.strip():
                    insert_paragraph_at_position(current_pos, resolucion.strip(), 
                                               alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, style_name="Body Text")
                    current_pos += 1
                """
                es_ultima = resolucion is resoluciones[-1]

                if isinstance(resolucion, dict):
                    clave = (resolucion.get('clave') or '').strip().rstrip('.')  # evita doble punto
                    texto = (resolucion.get('texto') or '').strip()
                elif isinstance(resolucion, str):
                    clave = ''
                    texto = resolucion.strip()
                else:
                    continue

                if not texto:
                    continue

                # 1) primer párrafo: clave (bold) + texto
                p = insert_paragraph_at_position(current_pos, "", style_name="Body Text")
                # sangría francesa: 1) sangría general; 2) primera línea negativa
                p.paragraph_format.left_indent = Cm(1.30)
                p.paragraph_format.first_line_indent = Cm(-1.30)
                p.paragraph_format.space_after = Pt(12)     # separación entre resoluciones
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                # aquí decides el espacio después
                p.paragraph_format.space_after = Pt(38 if es_ultima else 18)

                if clave:
                    r1 = p.add_run(f"{clave}. ")
                    p.add_run("\t")
                    r1.bold = False
                p.add_run(texto)
                current_pos += 1

                # 2) si el texto trae saltos de línea, crea párrafos de continuación (sin clave)
                #    que mantengan la misma sangría francesa y justificado
                #    (ej. notas largas en varias líneas)
                extra_parrafos = [t.strip() for t in texto.split("\n") if t.strip()]
                if len(extra_parrafos) > 1:
                    # ya imprimimos la primera parte arriba; añade el resto uno por párrafo
                    for cont in extra_parrafos[1:]:
                        p2 = insert_paragraph_at_position(current_pos, cont, style_name="Body Text")
                        p2.paragraph_format.left_indent = Cm(1.30)
                        p2.paragraph_format.first_line_indent = Cm(-1.30)
                        p2.paragraph_format.space_after = Pt(12)
                        p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        current_pos += 1
                
def test_ordenes_y_resoluciones():
    """
    Quick smoke test function to verify the functionality.
    """
    # Sample data
    orden_dia_json = [
        {
            "numero": 1,
            "titulo": "Perspectivas políticas y económicas",
            "descripcion": "Se discutieron las condiciones macro actuales.",
            "resoluciones": [
                {"clave": "Única", "tipo": "toma_nota", "texto": "Se toma nota de lo expuesto."}
            ]
        },
        {
            "numero": 2,
            "titulo": "Informe del Director General",
            "descripcion": "",
            "resoluciones": [
                {"clave": "II.1", "tipo": "aprobacion", "texto": "Se aprueba el informe del Director General."},
                {"clave": "II.2", "tipo": "instruccion", "texto": "Se instruye continuar con los análisis."}
            ]
        }
    ]
    
    # Test build_ordenes_con_resoluciones
    ordenes = build_ordenes_con_resoluciones(orden_dia_json)
    
    print("=== Test Results ===")
    print(f"Number of ordenes: {len(ordenes)}")
    
    for orden in ordenes:
        roman = to_roman(orden['numero'])
        print(f"{roman}. {orden['titulo']}")
        if orden['descripcion']:
            print(f"   Descripción: {orden['descripcion']}")
        print(f"   Resoluciones: {len(orden['resoluciones'])}")
        for res in orden['resoluciones']:
            if isinstance(res, dict):
                print(f"     - {res.get('clave', '')}: {res.get('texto', '')}")
    
    print("Test completed successfully!")


if __name__ == "__main__":
    test_ordenes_y_resoluciones()
