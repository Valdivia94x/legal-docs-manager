from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from django.conf import settings
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
import os
import re
from .models import ConvenioModificatorio
from num2words import num2words
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
def formatear_fecha(fecha):
    if not fecha:
        return "[FECHA]"
    fecha_str = fecha.strftime("%d de %B de %Y")
    meses = {
        'January': 'enero', 'February': 'febrero', 'March': 'marzo',
        'April': 'abril', 'May': 'mayo', 'June': 'junio',
        'July': 'julio', 'August': 'agosto', 'September': 'septiembre',
        'October': 'octubre', 'November': 'noviembre', 'December': 'diciembre'
    }
    for eng, esp in meses.items():
        fecha_str = fecha_str.replace(eng, esp)
    return fecha_str

def descargar_docx_convenio_modificatorio(request, pk):
    """
    Genera y descarga un archivo DOCX de Convenio Modificatorio con placeholders reemplazados
    """
    # Obtener el objeto ConvenioModificatorio
    convenio = get_object_or_404(ConvenioModificatorio, pk=pk, usuario=request.user)
    
    # Ruta al template DOCX
    template_path = os.path.join(
        settings.BASE_DIR.parent, 
        "DOCUMENTOS OLEA ABOGADOS", 
        "Convenios Modificatorios", 
        "Convenio Modificatorio PLACE.docx"
    )
    
    if not os.path.exists(template_path):
        return HttpResponse("Template file not found", status=404)
    
    try:
        # Cargar el documento template
        doc = Document(template_path)
        
        # Diccionario de reemplazos - mapeando campos del modelo a placeholders
        replacements = {
            # Datos Generales
            '{{fecha_convenio}}': formatear_fecha(convenio.fecha_convenio),
            '{{lugar_convenio}}': convenio.lugar_convenio or '',
            
            # Inversionista (F101)
            '{{inversionista_razon_social}}': convenio.inversionista_razon_social or '',
            '{{INVERSIONISTA_RAZON_SOCIAL}}': convenio.inversionista_razon_social.upper() or '',
            '{{inversionista_representante}}': convenio.inversionista_representante or '',
            '{{INVERSIONISTA_REPRESENTANTE}}': convenio.inversionista_representante.upper() or '',
            '{{inversionista_constitucion_escritura}}': convenio.inversionista_constitucion_escritura or '',
            '{{inversionista_constitucion_fecha}}': formatear_fecha(convenio.inversionista_constitucion_fecha),
            '{{inversionista_notario_constitucion}}': convenio.inversionista_notario_constitucion or '',
            '{{inversionista_registro_constitucion}}': convenio.inversionista_registro_constitucion or '',
            '{{inversionista_poder_escritura}}': convenio.inversionista_poder_escritura or '',
            '{{inversionista_poder_fecha}}': formatear_fecha(convenio.inversionista_poder_fecha),
            '{{inversionista_poder_notario}}': convenio.inversionista_poder_notario or '',
            
            # Antecedente: Contrato de Inversión Original
            '{{contrato_original_fecha}}': formatear_fecha(convenio.contrato_original_fecha),
            
            # Estudiante
            '{{estudiante_nombre}}': convenio.estudiante_nombre or '',
            '{{ESTUDIANTE_NOMBRE}}': convenio.estudiante_nombre.upper() or '',
            '{{estudiante_estado_civil}}': convenio.estudiante_estado_civil or '',
            '{{estudiante_nacionalidad}}': convenio.estudiante_nacionalidad or '',
            '{{estudiante_ocupacion}}': convenio.estudiante_ocupacion or '',
            '{{ESTUDIANTE_OCUPACION}}': convenio.estudiante_ocupacion.upper() or '',
            '{{estudiante_rfc}}': convenio.estudiante_rfc or '',
            '{{ESTUDIANTE_RFC}}': convenio.estudiante_rfc.upper() or '',
            '{{estudiante_curp}}': convenio.estudiante_curp or '',
            '{{ESTUDIANTE_CURP}}': convenio.estudiante_curp.upper() or '',
            '{{adeudo_principal_anterior}}': f"${convenio.adeudo_principal_anterior:,.2f}" if convenio.adeudo_principal_anterior else '',
            '{{adeudo_principal_anterior_texto}}': convenio.adeudo_principal_anterior_texto or '',
            '{{credito_original}}': f"${convenio.credito_original:,.2f}" if convenio.credito_original else '',
            '{{credito_original_texto}}': num2words(convenio.credito_original, lang='es'),
            
            # Obligados Solidarios
            '{{luis_nombre}}': convenio.luis_nombre or '',
            '{{LUIS_NOMBRE}}': convenio.luis_nombre.upper() or '',
            '{{luis_estado_civil}}': convenio.luis_estado_civil or '',
            '{{luis_ocupacion}}': convenio.luis_ocupacion or '',
            '{{luis_rfc}}': convenio.luis_rfc or '',
            '{{LUIS_RFC}}': convenio.luis_rfc.upper() or '',
            '{{luis_curp}}': convenio.luis_curp or '',
            '{{LUIS_CURP}}': convenio.luis_curp.upper() or '',
            
            '{{lizette_nombre}}': convenio.lizette_nombre or '',
            '{{LIZETTE_NOMBRE}}': convenio.lizette_nombre.upper() or '',
            '{{lizette_estado_civil}}': convenio.lizette_estado_civil or '',
            '{{lizette_ocupacion}}': convenio.lizette_ocupacion or '',
            '{{lizette_rfc}}': convenio.lizette_rfc or '',
            '{{lizette_curp}}': convenio.lizette_curp or '',
            
            # Aumento de la Inversión (Inversión II)
            '{{aumento_monto}}': f"${convenio.aumento_monto:,.2f}" if convenio.aumento_monto else '',
            '{{aumento_monto_texto}}': convenio.aumento_monto_texto or '',
            
            # Reconocimiento y Pagaré II
            '{{inversion_total}}': f"${convenio.inversion_total:,.2f}" if convenio.inversion_total else '',
            '{{inversion_total_texto}}': convenio.inversion_total_texto or '',
            '{{adeudo_actualizado}}': f"${convenio.adeudo_actualizado:,.2f}" if convenio.adeudo_actualizado else '',
            '{{adeudo_actualizado_texto}}': num2words(convenio.adeudo_actualizado, lang='es'),
            '{{pagare_ii_fecha}}': formatear_fecha(convenio.pagare_ii_fecha),
            '{{pagare_i_devuelto}}': 'Sí' if convenio.pagare_i_devuelto else 'No',
            
            # Periodo de Disposición
            '{{dispo_periodo_inicio}}': formatear_fecha(convenio.dispo_periodo_inicio),
            '{{dispo_periodo_fin}}': formatear_fecha(convenio.dispo_periodo_fin),
            
            # Términos Financieros
            '{{cat_anual}}': f"{convenio.cat_anual}%" if convenio.cat_anual else '',
            '{{tasa_interes_mensual}}': f"{convenio.tasa_interes_mensual}%" if convenio.tasa_interes_mensual else '',
            '{{tasa_moratoria_mensual}}': f"{convenio.tasa_moratoria_mensual}%" if convenio.tasa_moratoria_mensual else '',
            
            # Pagos en Etapa de Estudios
            '{{pagos_estudios_num1}}': str(convenio.pagos_estudios_num1) if convenio.pagos_estudios_num1 else '',
            '{{pagos_estudios_imp1}}': f"${convenio.pagos_estudios_imp1:,.2f}" if convenio.pagos_estudios_imp1 else '',
            '{{pagos_estudios_num2}}': str(convenio.pagos_estudios_num2) if convenio.pagos_estudios_num2 else '',
            '{{pagos_estudios_num2_texto}}': num2words(convenio.pagos_estudios_num2, lang='es'),
            '{{pagos_estudios_imp2}}': f"${convenio.pagos_estudios_imp2:,.2f}" if convenio.pagos_estudios_imp2 else '',
            
            # Pagos en Etapa de Egreso
            '{{pagos_egreso_num}}': str(convenio.pagos_egreso_num) if convenio.pagos_egreso_num else '',
            '{{PAGOS_EGRESO_NUM_TEXTO}}': num2words(convenio.pagos_egreso_num, lang='es'),
            '{{pagos_egreso_imp}}': f"${convenio.pagos_egreso_imp:,.2f}" if convenio.pagos_egreso_imp else '',
            '{{pago_egreso_ultimo}}': f"${convenio.pago_egreso_ultimo:,.2f}" if convenio.pago_egreso_ultimo else '',
            
            # Cuenta para pagos
            '{{cuenta_numero}}': convenio.cuenta_numero or '',
            '{{cuenta_banco}}': convenio.cuenta_banco or '',
            '{{cuenta_titular}}': convenio.cuenta_titular or '',
            '{{cuenta_clabe}}': convenio.cuenta_clabe or '',
            
            # Documentos a entregar (Cláusula Novena)
            '{{entrego_acta_nacimiento}}': 'Sí' if convenio.entrego_acta_nacimiento else 'No',
            '{{entrego_identificacion_oficial}}': 'Sí' if convenio.entrego_identificacion_oficial else 'No',
            '{{entrego_reporte_buro}}': 'Sí' if convenio.entrego_reporte_buro else 'No',
            '{{entrego_rfc}}': 'Sí' if convenio.entrego_rfc else 'No',
            '{{entrego_curp}}': 'Sí' if convenio.entrego_curp else 'No',
            '{{entrego_comprobante_domicilio}}': 'Sí' if convenio.entrego_comprobante_domicilio else 'No',
            '{{entrego_comprobante_ingresos}}': 'Sí' if convenio.entrego_comprobante_ingresos else 'No',
            
            # Beneficio por Cursos de Capacitación
            '{{cursos_requeridos}}': str(convenio.cursos_requeridos) if convenio.cursos_requeridos else '',
            '{{beneficio_descripcion}}': convenio.beneficio_descripcion or '',
            
            # Confidencialidad
            '{{confidencialidad_years}}': str(convenio.confidencialidad_years) if convenio.confidencialidad_years else '',
        }
        
        # Función auxiliar para reemplazar texto completo en párrafos
        def replace_in_paragraph(paragraph, replacements):
            # Obtener todo el texto del párrafo
            full_text = paragraph.text
            
            # Verificar si hay placeholders en el párrafo
            has_placeholder = False
            for placeholder in replacements.keys():
                if placeholder in full_text:
                    has_placeholder = True
                    break
            
            if not has_placeholder:
                return
            
            # Realizar reemplazos normales
            new_text = full_text
            for placeholder, value in replacements.items():
                if placeholder in new_text:
                    new_text = new_text.replace(placeholder, str(value))
            
            # Si el texto cambió, reemplazar todo el párrafo
            if new_text != full_text:
                # Limpiar todos los runs existentes
                for run in paragraph.runs:
                    run.text = ""
                
                # Agregar el nuevo texto al primer run
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                else:
                    paragraph.add_run(new_text)
        
        # Reemplazar placeholders en párrafos
        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph, replacements)
        
        # Reemplazar placeholders en tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph, replacements)
        
        # Reemplazar placeholders en headers y footers
        for section in doc.sections:
            # Header
            if section.header:
                for paragraph in section.header.paragraphs:
                    replace_in_paragraph(paragraph, replacements)
            
            # Footer
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    replace_in_paragraph(paragraph, replacements)
        
        # Preparar la respuesta HTTP
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        # Nombre del archivo de descarga
        filename = f"Convenio_Modificatorio_{convenio.estudiante_nombre.replace(' ', '_') if convenio.estudiante_nombre else 'documento'}_{convenio.pk}.docx"
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        # Guardar el documento en la respuesta
        doc.save(response)
        
        return response
        
    except Exception as e:
        return HttpResponse(f"Error generating document: {str(e)}", status=500)
