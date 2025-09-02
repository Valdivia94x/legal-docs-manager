import json
import os
from io import BytesIO
from .estatutos_sociedad import EstatutosSociedad
from .models import ActaAsamblea, ActaSesionConsejo, Pagare, ContratoCredito, ContratoPrendaAcciones, ConvenioModificatorio
from .forms_estatutos_fixed import EstatutosSociedadForm
from .docx_estatutos_generator import descargar_docx_estatutos_sociedad
from .docx_convenio_generator import descargar_docx_convenio_modificatorio
from .docx_prenda_generator import generar_docx_prenda
from django.http import HttpResponse
from django.shortcuts import get_object_or_404, redirect
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.contrib.auth import get_user_model
from django.conf import settings

# ReportLab imports
from reportlab.lib.pagesizes import A4, letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm, mm
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    HRFlowable,
    PageBreak,
)
from reportlab.lib import colors
from num2words import num2words

# DOCX imports
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# DOC imports (for older .doc files)
try:
    import docx2txt
    DOC_AVAILABLE = True
except ImportError:
    DOC_AVAILABLE = False

from .docx_blocks import build_ordenes_con_resoluciones, inject_ordenes_y_resoluciones


def to_roman(num):
    """Convierte un número entero a numeración romana"""
    val = [1000, 900, 500, 400, 100, 90, 50, 40, 10, 9, 5, 4, 1]
    syms = ["M", "CM", "D", "CD", "C", "XC", "L", "XL", "X", "IX", "V", "IV", "I"]
    roman = ''
    i = 0
    while num > 0:
        for _ in range(num // val[i]):
            roman += syms[i]
            num -= val[i]
        i += 1
    return roman


def generar_pdf_acta_consejo(acta):
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=2*cm,
        leftMargin=2*cm,
        topMargin=2*cm,
        bottomMargin=2*cm
    )

    # Definir estilos
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle', parent=styles['Heading1'], alignment=TA_CENTER,
        fontSize=14, fontName='Times-Bold', spaceBefore=6, spaceAfter=6
    )
    normal_style = ParagraphStyle(
        'CustomNormal', parent=styles['Normal'], alignment=TA_JUSTIFY,
        fontSize=11, fontName='Times-Roman', leading=16, spaceAfter=12
    )
    section_title_style = ParagraphStyle(
        'SectionTitle', parent=styles['Heading2'], alignment=TA_CENTER,
        fontSize=12, fontName='Times-Bold', spaceBefore=12, spaceAfter=12
    )
    agenda_item_style = ParagraphStyle(
        'AgendaItem', parent=styles['Normal'], alignment=TA_LEFT,
        fontSize=11, fontName='Times-Roman',
        leftIndent=1.5*cm, firstLineIndent=-0.5*cm,
        spaceBefore=8, spaceAfter=8
    )
    resolution_title_style = ParagraphStyle(
        'ResolutionTitle', parent=styles['Normal'], alignment=TA_CENTER,
        fontSize=11, fontName='Times-Bold', spaceBefore=12, spaceAfter=6
    )
    resolution_text_style = ParagraphStyle(
        'ResolutionText', parent=styles['Normal'], alignment=TA_JUSTIFY,
        fontSize=11, fontName='Times-Roman', leftIndent=2*cm, spaceAfter=8
    )

    story = []

    # Cabecera
    story.append(Paragraph("SESIÓN DEL CONSEJO DE ADMINISTRACIÓN DE", title_style))
    story.append(Paragraph(acta.razon_social.upper(), title_style))

    # Fecha
    meses = {1: "enero", 2: "febrero", 3: "marzo", 4: "abril", 5: "mayo", 6: "junio",
             7: "julio", 8: "agosto", 9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre"}
    fecha = f"{acta.fecha.day} de {meses[acta.fecha.month]} de {acta.fecha.year}"
    story.append(Paragraph(f"DEL DÍA {fecha.upper()}", title_style))
    story.append(Spacer(1, 12))

    # Apertura
    hora_inicio = acta.hora_inicio.strftime("%H:%M") if acta.hora_inicio else "[HORA]"
    lugar = acta.lugar or "[LUGAR DE LA SESIÓN]"
    apertura = (
        f"En la Ciudad de México, México siendo las {hora_inicio} horas del día {fecha}, "
        f"se reunieron dentro del domicilio social de {acta.razon_social} (en lo sucesivo, "
        f"indistintamente como “la Sociedad”) en su domicilio, sito en {lugar}, o de manera remota, "
        f"los señores Consejeros que aparecen en la lista de asistencia que se adjunta como Anexo Uno, "
        f"con el propósito de celebrar una sesión del Consejo de Administración, a la cual fueron convocados "
        f"según consta en la convocatoria que se adjunta como Anexo Dos."
    )
    story.append(Paragraph(apertura, normal_style))

    # Invitados
    if acta.invitados_json:
        try:
            invitados = acta.invitados_json
            if isinstance(invitados, str):
                invitados = json.loads(invitados)
            nombres = [i.get('nombre') for i in invitados if isinstance(i, dict) and i.get('nombre')]
            if nombres:
                story.append(Paragraph(
                    f"Se hace constar la presencia de {', '.join(nombres)}, quienes comparecen como invitados.",
                    normal_style
                ))
        except:
            pass

    # Presidencia e instalación
    pres = (
        f"De conformidad con la Cláusula Vigésima Quinta de los estatutos, presidió la sesión {acta.presidente}, Presidente del Consejo, "
        f"y {acta.secretario}, Secretario del Consejo."
    )
    story.append(Paragraph(pres, normal_style))
    story.append(Paragraph(
        "El Presidente declaró legalmente instalada la sesión al encontrarse presente la mayoría de los miembros, "
        "de conformidad con la Cláusula Vigésima Octava.",
        normal_style
    ))

    # Orden del Día
    story.append(Spacer(1, 12))
    story.append(Paragraph("ORDEN DEL DÍA", section_title_style))
    if acta.orden_dia_json:
        try:
            items = acta.orden_dia_json
            if isinstance(items, str):
                items = json.loads(items)
            for idx, itm in enumerate(items, 1):
                roman = to_roman(idx)
                titulo = itm.get('titulo') if isinstance(itm, dict) else str(itm)
                story.append(Paragraph(f"{roman}. {titulo}", agenda_item_style))
        except Exception as e:
            story.append(Paragraph(f"Error al procesar orden del día: {e}", normal_style))

    # Transición
    story.append(Spacer(1, 12))
    story.append(Paragraph(
        "Los Consejeros, tras la lectura del orden del día, procedieron a discutir los asuntos en él contenidos:",
        normal_style
    ))

    # Desarrollo y resoluciones
    if acta.resoluciones_json:
        try:
            res = acta.resoluciones_json
            if isinstance(res, str):
                res = json.loads(res)
            story.append(Spacer(1, 12))
            story.append(Paragraph("DESARROLLO DE LA SESIÓN", section_title_style))
            for idx, r in enumerate(res, 1):
                roman = to_roman(idx)
                titulo = r.get('titulo', f"Punto {idx}") if isinstance(r, dict) else str(r)
                story.append(Paragraph(f"{roman}. {titulo}", agenda_item_style))
                desc = r.get('descripcion', "Se discutió el punto correspondiente.")
                story.append(Paragraph(desc, normal_style))
                story.append(Paragraph("R E S O L U C I Ó N", resolution_title_style))
                clave = r.get('clave', 'Única')
                texto = r.get('texto', '')
                story.append(Paragraph(f"<b>{clave}.</b> \"{texto}\"", resolution_text_style))
        except Exception as e:
            story.append(Paragraph(f"Error al procesar resoluciones: {e}", normal_style))

    # Cierre
    hora_fin = acta.hora_cierre.strftime("%H:%M") if acta.hora_cierre else "[HORA]"
    cierre = (
        f"No habiendo otro asunto, se suspendió la sesión para redacción del acta, la cual, "
        f"una vez leída por el Secretario, fue aprobada y concluyó a las {hora_fin} horas."
    )
    story.append(Spacer(1, 12))
    story.append(Paragraph(cierre, normal_style))

    # Firmas
    story.append(Spacer(1, 48))
    firma_data = [
        [HRFlowable(width=6*cm, thickness=1), HRFlowable(width=6*cm, thickness=1)],
        [acta.presidente, acta.secretario]
    ]
    firma_tbl = Table(firma_data, colWidths=[7*cm, 7*cm], hAlign='CENTER')
    firma_tbl.setStyle(TableStyle([
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('FONTNAME', (1,1), (-1,-1), 'Times-Roman'),
        ('FONTSIZE', (0,0), (-1,-1), 11),
        ('TOPPADDING', (1,1), (-1,-1), 6)
    ]))
    story.append(firma_tbl)

    # Generar PDF
    doc.build(story)
    pdf = buffer.getvalue()
    buffer.close()
    return pdf


def generar_pdf_pagare(pagare):
    """Genera PDF para Pagaré con formato legal profesional"""
    buffer = BytesIO()
    
    # Configurar documento con márgenes específicos
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=2.5*cm,
        leftMargin=2.5*cm,
        topMargin=2*cm,
        bottomMargin=2*cm
    )
    
    # Crear estilos personalizados
    styles = getSampleStyleSheet()
    
    # Estilo para título principal (centrado, negrita, 14pt)
    title_style = ParagraphStyle(
        'PagareTitle',
        parent=styles['Heading1'],
        fontSize=14,
        spaceAfter=20,
        alignment=TA_CENTER,
        fontName='Times-Bold'
    )
    
    # Estilo para subtítulos (centrado, negrita, 12pt)
    subtitle_style = ParagraphStyle(
        'PagareSubtitle',
        parent=styles['Heading2'],
        fontSize=12,
        spaceAfter=12,
        alignment=TA_CENTER,
        fontName='Times-Bold'
    )
    
    # Estilo para párrafos justificados (11pt)
    justified_style = ParagraphStyle(
        'PagareJustified',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=12,
        alignment=TA_JUSTIFY,
        fontName='Times-Roman',
        leading=14
    )
    
    # Estilo para texto centrado
    center_style = ParagraphStyle(
        'PagareCenter',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=12,
        alignment=TA_CENTER,
        fontName='Times-Roman'
    )
    
    story = []
    
    # ENCABEZADO
    story.append(Spacer(1, 20))
    story.append(Paragraph("PAGARÉ", title_style))
    story.append(Spacer(1, 20))
    
    # Lugar y fecha de emisión
    lugar_fecha = f"{pagare.lugar_emision}, {pagare.fecha_emision.strftime('%d de %B de %Y')}"
    # Convertir mes a español
    meses = {
        'January': 'enero', 'February': 'febrero', 'March': 'marzo',
        'April': 'abril', 'May': 'mayo', 'June': 'junio',
        'July': 'julio', 'August': 'agosto', 'September': 'septiembre',
        'October': 'octubre', 'November': 'noviembre', 'December': 'diciembre'
    }
    for eng, esp in meses.items():
        lugar_fecha = lugar_fecha.replace(eng, esp)
    
    story.append(Paragraph(lugar_fecha, center_style))
    story.append(Spacer(1, 20))
    
    # CUERPO PRINCIPAL DEL PAGARÉ
    # Párrafo de reconocimiento de deuda
    monto_texto = f"${pagare.monto_numeric:,.2f} {pagare.moneda}"
    parrafo_principal = f"""Por medio del presente pagaré, yo <b>{pagare.deudor_nombre}</b>, con domicilio en {pagare.deudor_domicilio}, reconozco deber y me obligo a pagar incondicionalmente a la orden de <b>{pagare.acreedor_nombre}</b>, con domicilio en {pagare.acreedor_domicilio}, la cantidad de <b>{monto_texto}</b> ({pagare.monto_literal}), por concepto de {pagare.concepto}."""
    
    story.append(Paragraph(parrafo_principal, justified_style))
    story.append(Spacer(1, 15))
    
    # CONDICIONES DE PAGO
    story.append(Paragraph("CONDICIONES DE PAGO", subtitle_style))
    
    if pagare.tipo_pago == 'unico':
        condiciones_pago = f"El presente pagaré será pagado en un solo pago."
    else:
        condiciones_pago = f"El presente pagaré será pagado en {pagare.num_pagos} parcialidades {pagare.get_periodicidad_display().lower()}s."
    
    if pagare.lugar_pago:
        condiciones_pago += f" El pago deberá realizarse en {pagare.lugar_pago}."
    
    if pagare.forma_pago:
        condiciones_pago += f" La forma de pago será mediante {pagare.forma_pago}."
    
    story.append(Paragraph(condiciones_pago, justified_style))
    story.append(Spacer(1, 15))
    
    # INTERESES
    story.append(Paragraph("INTERESES Y CARGOS", subtitle_style))
    
    intereses_texto = f"""El presente pagaré devengará intereses ordinarios a una tasa del {pagare.tasa_interes_ordinario}% anual sobre saldos insolutos, calculados sobre una base de {pagare.base_intereses} días. En caso de mora, se aplicará una tasa de interés moratorio del {pagare.tasa_interes_moratorio}% anual."""
    
    if pagare.gastos_admon > 0:
        intereses_texto += f" Adicionalmente, se aplicarán gastos de administración por ${pagare.gastos_admon:,.2f}."
    
    story.append(Paragraph(intereses_texto, justified_style))
    story.append(Spacer(1, 15))
    
    # PREPAGO (si aplica)
    if pagare.prepago_permitido:
        story.append(Paragraph("PREPAGO", subtitle_style))
        prepago_texto = f"El deudor podrá realizar prepagos"
        if pagare.dias_aviso_prepago:
            prepago_texto += f" con {pagare.dias_aviso_prepago} días de aviso previo"
        if pagare.condicion_prepago:
            prepago_texto += f", {pagare.condicion_prepago}"
        prepago_texto += "."
        story.append(Paragraph(prepago_texto, justified_style))
        story.append(Spacer(1, 15))
    
    # GARANTÍAS (si aplica)
    if pagare.tiene_garantia:
        story.append(Paragraph("GARANTÍAS Y AVALES", subtitle_style))
        garantia_texto = f"El presente pagaré se encuentra garantizado mediante: {pagare.descripcion_garantia}"
        if pagare.aval_nombre:
            garantia_texto += f" Actúa como aval el Sr./Sra. {pagare.aval_nombre}, con domicilio en {pagare.aval_domicilio}."
        story.append(Paragraph(garantia_texto, justified_style))
        story.append(Spacer(1, 15))
    
    # TABLA DE AMORTIZACIÓN (si existe)
    if pagare.tabla_amortizacion:
        story.append(Paragraph("TABLA DE AMORTIZACIÓN", subtitle_style))
        
        # Crear tabla con los pagos
        tabla_data = [['No.', 'Fecha', 'Capital', 'Intereses', 'Total', 'Saldo']]
        
        try:
            cuotas = pagare.obtener_cuotas()
            for cuota in cuotas[:10]:  # Mostrar máximo 10 cuotas en PDF
                tabla_data.append([
                    str(cuota['numero']),
                    cuota['fecha'][:10] if isinstance(cuota['fecha'], str) else str(cuota['fecha']),
                    f"${cuota['capital']:,.2f}",
                    f"${cuota['interes']:,.2f}",
                    f"${cuota['total']:,.2f}",
                    f"${cuota['saldo']:,.2f}"
                ])
        except:
            tabla_data.append(['1', '[FECHA]', '[CAPITAL]', '[INTERESES]', '[TOTAL]', '[SALDO]'])
        
        tabla = Table(tabla_data, colWidths=[1*cm, 2.5*cm, 2.5*cm, 2.5*cm, 2.5*cm, 2.5*cm])
        tabla.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Times-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(tabla)
        story.append(Spacer(1, 15))
    
    # CLÁUSULAS LEGALES
    story.append(Paragraph("DISPOSICIONES LEGALES", subtitle_style))
    
    clausulas_texto = f"""El presente pagaré se rige por las siguientes disposiciones: {pagare.eventos_incumplimiento} {pagare.clausula_aceleracion} Para la interpretación y cumplimiento del presente pagaré, las partes se someten a la jurisdicción de {pagare.jurisdiccion}, renunciando al fuero que por razón de su domicilio presente o futuro pudiera corresponderles. El presente pagaré se rige por {pagare.ley_aplicable}."""
    
    story.append(Paragraph(clausulas_texto, justified_style))
    story.append(Spacer(1, 20))
    
    # ANEXOS (si existen)
    if pagare.anexos:
        story.append(Paragraph("ANEXOS", subtitle_style))
        story.append(Paragraph(pagare.anexos, justified_style))
        story.append(Spacer(1, 15))
    
    # FIRMAS
    story.append(Spacer(1, 30))
    story.append(Paragraph("FIRMAS", subtitle_style))
    story.append(Spacer(1, 20))
    
    # Crear tabla de firmas
    firmas_data = [
        ["_" * 40, "_" * 40],
        [pagare.deudor_nombre, pagare.acreedor_nombre],
        ["DEUDOR", "ACREEDOR"]
    ]
    
    if pagare.aval_nombre:
        firmas_data.append(["", ""])
        firmas_data.append(["_" * 40, ""])
        firmas_data.append([pagare.aval_nombre, ""])
        firmas_data.append(["AVAL", ""])
    
    firmas_table = Table(firmas_data, colWidths=[8*cm, 8*cm])
    firmas_table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('FONTNAME', (0, 0), (-1, -1), 'Times-Roman'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('TOPPADDING', (0, 0), (-1, -1), 12),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
    ]))
    
    story.append(firmas_table)
    story.append(Spacer(1, 20))
    
    # PIE DE DOCUMENTO
    story.append(Paragraph(f"Documento de {pagare.num_paginas} página(s)", center_style))
    
    # Construir PDF
    doc.build(story)
    pdf = buffer.getvalue()
    buffer.close()
    return pdf


@login_required
def descargar_pdf_pagare(request, pk):
    """Vista para descargar PDF de Pagaré"""
    pagare = get_object_or_404(Pagare, pk=pk, usuario=request.user)
    try:
        pdf = generar_pdf_pagare(pagare)
        response = HttpResponse(pdf, content_type='application/pdf')
        filename = f'pagare_{pagare.pk}_{pagare.fecha_emision.strftime("%Y%m%d")}.pdf'
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response
    except Exception as e:
        messages.error(request, f'Error al generar PDF: {e}')
        return redirect('documentos:detalle_pagare', pk=pk)


def replace_in_xml_textnodes(part, mapping):
    """
    Reemplaza en todos los nodos <w:t> de un part (documento, header, footer, etc.)
    """
    # Selecciona todos los nodos de texto
    for t in part.element.xpath(".//w:t"):
        if t.text:
            for k, v in mapping.items():
                if k in t.text:
                    t.text = t.text.replace(k, v)


def generar_docx_acta_consejo(acta):
    """Genera un documento DOCX basado en template con reemplazos de datos del acta"""
    if not DOCX_AVAILABLE:
        raise ImportError("La librería python-docx no está disponible")
    
    # Ruta del template
    template_path = os.path.join(
        settings.BASE_DIR.parent, 
        "DOCUMENTOS OLEA ABOGADOS", 
        "Actas Sesiones de Consejo", 
        "Acta de Sesion de Consejo de Administración PLACE.docx"
    )
    
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template no encontrado en: {template_path}")
    
    # Cargar el template
    doc = Document(template_path)
    
    # Preparar datos para reemplazo
    fecha_formateada = acta.fecha.strftime("%d de %B de %Y") if acta.fecha else "[FECHA]"
    # Convertir mes a español
    meses = {
        'January': 'enero', 'February': 'febrero', 'March': 'marzo',
        'April': 'abril', 'May': 'mayo', 'June': 'junio',
        'July': 'julio', 'August': 'agosto', 'September': 'septiembre',
        'October': 'octubre', 'November': 'noviembre', 'December': 'diciembre'
    }
    for eng, esp in meses.items():
        fecha_formateada = fecha_formateada.replace(eng, esp)
    
    def formatear_invitado(inv):
        """
        Devuelve 'el ingeniero Juan Pérez', 'la licenciada Ana Ruiz' o solo 'Juan Pérez (Contador Público)'.
        Si no hay cargo/título, solo usa el nombre.
        """
        nombre = inv.get("nombre", "").strip()
        cargo  = (inv.get("cargo") or "").strip().lower()

        # Heurística simple para artículos (ajústala a tu dataset)
        femeninos = {"licenciada", "ingeniera", "abogada", "contadora", "doctora", "arquitecta"}
        masculinos = {"licenciado", "ingeniero", "abogado", "contador", "doctor", "arquitecto"}

        if cargo in femeninos:
            return f"la {cargo} {nombre}"
        elif cargo in masculinos:
            return f"el {cargo} {nombre}"
        elif cargo:
            # Cargo sin género marcado: lo dejamos pospuesto y entre paréntesis
            return f"{nombre} ({inv.get('cargo')})"
        else:
            return nombre or "Invitado"

    def unir_conyuncion(nombres):
        """
        Une con comas y 'y' antes del último:
        ['A'] -> 'A'
        ['A','B'] -> 'A y B'
        ['A','B','C'] -> 'A, B y C'
        """
        if not nombres:
            return ""
        if len(nombres) == 1:
            return nombres[0]
        if len(nombres) == 2:
            return f"{nombres[0]} y {nombres[1]}"
        return f"{', '.join(nombres[:-1])} y {nombres[-1]}"

    def oracion_invitados(invitados):
        """
        Genera la oración completa, variando singular/plural y posición de 'asimismo'.
        Ejemplos:
          0 -> '' (o una línea neutra si prefieres)
          1 -> 'Se hace constar la presencia a la sesión de la licenciada Ana Ruiz, quien comparece en calidad de invitada.'
          2 -> 'Se hace constar la presencia a la sesión del ingeniero Juan Pérez y la licenciada Ana Ruiz, quienes comparecen en calidad de invitados.'
          3+-> 'Se hace constar la presencia a la sesión de Juan, Ana y Carlos, quienes comparecen en calidad de invitados.'
        """
        lista = [formatear_invitado(i) for i in (invitados or []) if i.get("nombre")]

        if not lista:
            return ""  # o: "Se hace constar que no hubo invitados a la sesión."

        # Artículo inicial: si el primero ya trae 'el/la', usa 'de'; si no, también 'de'
        sujetos = unir_conyuncion(lista)

        if len(lista) == 1:
            # Ajusta 'invitado/a' según heurística del formateo (muy simple):
            texto = lista[0]
            invitadx = "invitada" if texto.startswith("la ") else "invitado"
            return f"Se hace constar la presencia a la sesión de {texto}, quien comparece en calidad de {invitadx}."
        else:
            return f"Se hace constar la presencia a la sesión de {sujetos}, quienes comparecen en calidad de invitados."

    # Obtener invitados con debugging
    invitados = []
    invitados_data_raw = []
    print(f"DEBUG: acta.invitados_json = {acta.invitados_json}")
    print(f"DEBUG: type(acta.invitados_json) = {type(acta.invitados_json)}")
    
    if acta.invitados_json:
        try:
            invitados_data = acta.invitados_json
            if isinstance(invitados_data, str):
                invitados_data = json.loads(invitados_data)
            
            print(f"DEBUG: invitados_data after parsing = {invitados_data}")
            print(f"DEBUG: type(invitados_data) = {type(invitados_data)}")
            
            if isinstance(invitados_data, list):
                invitados_data_raw = invitados_data  # Keep raw data for oracion_invitados
                for inv in invitados_data:
                    if isinstance(inv, dict) and 'nombre' in inv:
                        nombre = inv.get('nombre', '').strip()
                        if nombre:
                            invitados.append(nombre)
                            print(f"DEBUG: Added invitado: {nombre}")
            
            print(f"DEBUG: Final invitados list = {invitados}")
        except Exception as e:
            print(f"DEBUG: Error processing invitados_json: {e}")
            pass
    
    # Generate proper invitados sentence
    oracion_completa_invitados = oracion_invitados(invitados_data_raw)
    
    # Crear diccionario de reemplazos con manejo de campos vacíos y verificación de atributos
    replacements = {
        "{{SOCIEDAD}}": acta.razon_social or "[SOCIEDAD]",
        "{{RAZON_SOCIAL}}": acta.razon_social or "[RAZON_SOCIAL]",
        "{{ABREVIATURA SOCIEDAD}}": (getattr(acta, 'abreviatura_sociedad', '').strip().upper() if hasattr(acta, 'abreviatura_sociedad') and getattr(acta, 'abreviatura_sociedad', '').strip() else 
                                    (acta.razon_social[:3].upper() if acta.razon_social else "CAS")),
        "{{FECHA}}": fecha_formateada,
        "{{HORA}}": acta.hora_inicio.strftime("%H:%M") if acta.hora_inicio else "[HORA]",
        "{{CIUDAD}}": (getattr(acta, 'ciudad', '').strip() if hasattr(acta, 'ciudad') and getattr(acta, 'ciudad', '').strip() else "Ciudad de México"),
        "{{LUGAR}}": acta.lugar or "[LUGAR]",
        "{{DOMICILIO_SOCIAL}}": acta.lugar or "[DOMICILIO SOCIAL]",
        "{{PLATAFORMA_REMOTA}}": (getattr(acta, 'plataforma_remota', '').strip() if hasattr(acta, 'plataforma_remota') and getattr(acta, 'plataforma_remota', '').strip() else "Zoom"),
        "{{INVITADO_1}}": invitados[0] if len(invitados) > 0 else "{{INVITADO_1}}",
        "{{INVITADO_2}}": invitados[1] if len(invitados) > 1 else "{{INVITADO_2}}",
        "{{INVITADO_3}}": invitados[2] if len(invitados) > 2 else "{{INVITADO_3}}",
        "{{ORACION_INVITADOS}}": oracion_completa_invitados,
        "{{PRESIDENTE}}": acta.presidente or "[PRESIDENTE]",
        "{{SECRETARIO}}": acta.secretario or "[SECRETARIO]",
        "{{HORA_CIERRE}}": acta.hora_cierre.strftime("%H:%M") if acta.hora_cierre else "[HORA CIERRE]"
    }
    
    # Agregar lógica condicional para CONVOCATORIA_REALIZADA
    if acta.convocatoria_realizada:
        replacements["{{CONVOCATORIA_TEXTO}}"] = (
            "a la cual fueron debidamente convocados según consta en la convocatoria "
            "que se adjunta a la presente como Anexo Dos."
        )
    else:
        replacements["{{CONVOCATORIA_TEXTO}}"] = (
            "a la cual asistieron sin necesidad de convocatoria formal, dejando constancia de lo anterior."
        )
    
    # Formatear resoluciones para el campo {{RESOLUCIONES}}
    resoluciones_texto = ""
    if acta.resoluciones_json:
        try:
            resoluciones_data = acta.resoluciones_json
            if isinstance(resoluciones_data, str):
                resoluciones_data = json.loads(resoluciones_data)
            
            if isinstance(resoluciones_data, list):
                resoluciones_lines = []
                for i, resolucion in enumerate(resoluciones_data, 1):
                    if isinstance(resolucion, dict):
                        clave = resolucion.get('clave', f'R{i}')
                        texto = resolucion.get('texto', '')
                        if texto:
                            resoluciones_lines.append(f"{texto}")
                    else:
                        resoluciones_lines.append(f"{str(resolucion)}")
                
                resoluciones_texto = '\n\n'.join(resoluciones_lines)
            else:
                resoluciones_texto = str(resoluciones_data)
        except Exception as e:
            print(f"DEBUG: Error processing resoluciones_json: {e}")
            resoluciones_texto = "[Error al procesar resoluciones]"
    
    if not resoluciones_texto:
        resoluciones_texto = "[No hay resoluciones registradas]"
    
    replacements["{{RESOLUCIONES}}"] = resoluciones_texto
    
    # Enhanced debugging for placeholder detection
    placeholder_found = False
    placeholder_locations = []
    
    # Check main document paragraphs
    for i, paragraph in enumerate(doc.paragraphs):
        if "{{ORDENES_Y_RESOLUCIONES}}" in paragraph.text:
            placeholder_found = True
            placeholder_locations.append(f"Main doc paragraph {i}: '{paragraph.text[:100]}...'")
    
    # Check tables
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, paragraph in enumerate(cell.paragraphs):
                    if "{{ORDENES_Y_RESOLUCIONES}}" in paragraph.text:
                        placeholder_found = True
                        placeholder_locations.append(f"Table {t_idx}, Row {r_idx}, Cell {c_idx}, Para {p_idx}: '{paragraph.text[:100]}...'")
    
    print(f"DEBUG: Placeholder {{{{ORDENES_Y_RESOLUCIONES}}}} found: {placeholder_found}")
    if placeholder_locations:
        print("DEBUG: Placeholder locations:")
        for loc in placeholder_locations:
            print(f"  - {loc}")
    
    # Also check for similar placeholders that might be typos
    similar_placeholders = []
    all_text = ""
    for paragraph in doc.paragraphs:
        all_text += paragraph.text + " "
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    all_text += paragraph.text + " "
    
    # Look for any placeholder that contains "ORDEN" or "RESOLUCION"
    import re
    orden_matches = re.findall(r'\{\{[^}]*ORDEN[^}]*\}\}', all_text, re.IGNORECASE)
    resol_matches = re.findall(r'\{\{[^}]*RESOLUCI[^}]*\}\}', all_text, re.IGNORECASE)
    
    if orden_matches or resol_matches:
        print(f"DEBUG: Found similar placeholders - ORDEN: {orden_matches}, RESOLUCION: {resol_matches}")
    
    # Inject "Orden del Día + Resoluciones" section with proper formatting
    if placeholder_found:
        try:
            print(f"DEBUG: acta.orden_dia_json type: {type(acta.orden_dia_json)}")
            print(f"DEBUG: acta.orden_dia_json value: {acta.orden_dia_json}")
            print(f"DEBUG: acta.resoluciones_json type: {type(acta.resoluciones_json)}")
            print(f"DEBUG: acta.resoluciones_json value: {acta.resoluciones_json}")
            
            ordenes = build_ordenes_con_resoluciones(acta.orden_dia_json, acta.resoluciones_json)
            print(f"DEBUG: Built ordenes count: {len(ordenes) if ordenes else 0}")
            if ordenes:
                print(f"DEBUG: First orden: {ordenes[0]}")
            
            texto_constante = "Los señores consejeros después de escuchar el orden del día antes transcrito procedieron a discutir ampliamente todos y cada uno de los asuntos contenidos en el mismo, desahogándose de la siguiente manera:"
            
            inject_ordenes_y_resoluciones(
                doc,
                placeholder="{{ORDENES_Y_RESOLUCIONES}}",
                ordenes=ordenes,
                texto_constante=texto_constante
            )
            print("DEBUG: Successfully injected ORDENES_Y_RESOLUCIONES section with proper formatting")
        except ValueError as e:
            print(f"DEBUG: ValueError during injection: {e}")
            # Fallback to text replacement
            replacements["{{ORDENES_Y_RESOLUCIONES}}"] = "[Error: No se pudo inyectar el formato - placeholder no encontrado]"
        except Exception as e:
            print(f"DEBUG: Error injecting ORDENES_Y_RESOLUCIONES: {e}")
            import traceback
            traceback.print_exc()
            # Fallback to text replacement
            replacements["{{ORDENES_Y_RESOLUCIONES}}"] = "[Error al procesar orden del día y resoluciones]"
    else:
        print("DEBUG: Skipping injection - placeholder not found in template")
        # Add fallback replacement
        replacements["{{ORDENES_Y_RESOLUCIONES}}"] = "[ORDENES_Y_RESOLUCIONES placeholder not found in template]"
    
    # Debug: Print all replacements
    print("DEBUG: Replacements dictionary:")
    for key, value in replacements.items():
        print(f"  '{key}' -> '{value}'")
    
    print(f"DEBUG: acta.razon_social = '{acta.razon_social}'")
    print(f"DEBUG: acta.lugar = '{acta.lugar}'")
    
    # Test if the XML replacement function is working
    print("DEBUG: Testing XML replacement function...")
    
    # Aplicar reemplazos en documento principal
    replace_in_xml_textnodes(doc.part, replacements)
    
    # Función auxiliar para reemplazar texto completo en párrafos
    def replace_in_paragraph(paragraph, replacements):
        # Obtener todo el texto del párrafo
        full_text = paragraph.text
        
        # Debug: imprimir texto encontrado
        if full_text.strip():
            print(f"DEBUG: Párrafo encontrado: '{full_text}'")
        
        # Verificar si hay placeholders en el párrafo
        has_placeholder = False
        found_placeholders = []
        for placeholder in replacements.keys():
            if placeholder in full_text:
                has_placeholder = True
                found_placeholders.append(placeholder)
        
        if found_placeholders:
            print(f"DEBUG: Placeholders encontrados: {found_placeholders}")
        
        if not has_placeholder:
            return
        
        # Realizar reemplazos en el texto completo
        new_text = full_text
        for placeholder, value in replacements.items():
            if placeholder in new_text:
                print(f"DEBUG: Reemplazando '{placeholder}' con '{value}'")
                new_text = new_text.replace(placeholder, str(value))
        
        print(f"DEBUG: Texto original: '{full_text}'")
        print(f"DEBUG: Texto nuevo: '{new_text}'")
        
        # Si el texto cambió, reemplazar todo el párrafo
        if new_text != full_text:
            # Limpiar todos los runs existentes
            for run in paragraph.runs:
                run.text = ""
            
            # Agregar el nuevo texto al primer run
            if paragraph.runs:
                paragraph.runs[0].text = new_text
            else:
                paragraph.add_run(new_text)
            print(f"DEBUG: Párrafo actualizado exitosamente")
    
    # Also try direct text replacement in paragraphs as backup
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, replacements)
    
    # Aplicar reemplazos en headers y footers
    for section in doc.sections:
        for header in (section.header, section.footer):
            replace_in_xml_textnodes(header.part, replacements)
    
    # Guardar en memoria
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()


@login_required
def descargar_pdf_consejo(request, pk):
    acta = get_object_or_404(ActaSesionConsejo, pk=pk, usuario=request.user)
    try:
        pdf = generar_pdf_acta_consejo(acta)
        response = HttpResponse(pdf, content_type='application/pdf')
        filename = f'acta_consejo_{acta.pk}_{acta.fecha.strftime("%Y%m%d")}.pdf'
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response
    except Exception as e:
        messages.error(request, f'Error al generar PDF: {e}')
        return redirect('documentos:detalle_consejo', pk=pk)


@login_required
def descargar_docx_consejo(request, pk):
    """Vista para descargar DOCX de Acta de Sesión de Consejo"""
    acta = get_object_or_404(ActaSesionConsejo, pk=pk, usuario=request.user)
    try:
        docx_content = generar_docx_acta_consejo(acta)
        response = HttpResponse(
            docx_content, 
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        filename = f'acta_consejo_{acta.pk}_{acta.fecha.strftime("%Y%m%d")}.docx'
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response
    except Exception as e:
        messages.error(request, f'Error al generar DOCX: {e}')
        return redirect('documentos:detalle_consejo', pk=pk)

def generar_pdf_acta_asamblea(acta):
    # Crear buffer para el PDF
    buffer = BytesIO()
    
    # Configurar documento con márgenes específicos (3cm como en el HTML)
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=3*cm,
        leftMargin=3*cm,
        topMargin=2.5*cm,
        bottomMargin=2.5*cm
    )
    
    # Crear estilos personalizados que repliquen el HTML
    styles = getSampleStyleSheet()
    
    # Estilo para título principal (centrado, negrita, 11pt)
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=11,
        spaceAfter=12,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    # Estilo para fecha (centrado, normal, 11pt)
    date_style = ParagraphStyle(
        'CustomDate',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=12,
        alignment=TA_CENTER,
        fontName='Helvetica'
    )
    
    # Estilo para párrafos justificados (11pt, interlineado 150%)
    justified_style = ParagraphStyle(
        'CustomJustified',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=12,
        alignment=TA_JUSTIFY,
        fontName='Helvetica',
        leading=16.5  # 150% de 11pt
    )
    
    # Estilo para subtítulos centrados (negrita, 11pt)
    subtitle_style = ParagraphStyle(
        'CustomSubtitle',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=12,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    # Estilo para elementos del orden del día (con sangría)
    order_style = ParagraphStyle(
        'CustomOrder',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=12,
        alignment=TA_JUSTIFY,
        fontName='Helvetica',
        leftIndent=1*cm,
        firstLineIndent=-1*cm,
        leading=16.5
    )
    
    story = []
    
    # PÁGINA 1 - PORTADA
    # Espacios en blanco iniciales
    for _ in range(8):
        story.append(Spacer(1, 12))
    
    # Línea horizontal
    story.append(Spacer(1, 12))
    
    # Más espacios
    for _ in range(10):
        story.append(Spacer(1, 12))
    
    # Título principal
    tipo_asamblea_display = acta.get_tipo_asamblea_display().upper() if hasattr(acta, 'get_tipo_asamblea_display') else acta.tipo_asamblea.upper()
    titulo_principal = f"ASAMBLEA GENERAL {tipo_asamblea_display} DE ACCIONISTAS DE"
    story.append(Paragraph(titulo_principal, title_style))
    story.append(Paragraph(acta.razon_social.upper(), title_style))
    
    # Espacios
    for _ in range(8):
        story.append(Spacer(1, 12))
    
    # Fecha
    fecha_formateada = acta.fecha.strftime('%d de %B de %Y')
    # Convertir mes a español
    meses = {
        'January': 'enero', 'February': 'febrero', 'March': 'marzo',
        'April': 'abril', 'May': 'mayo', 'June': 'junio',
        'July': 'julio', 'August': 'agosto', 'September': 'septiembre',
        'October': 'octubre', 'November': 'noviembre', 'December': 'diciembre'
    }
    for eng, esp in meses.items():
        fecha_formateada = fecha_formateada.replace(eng, esp)
    
    story.append(Paragraph(f"De fecha {fecha_formateada}", date_style))
    
    # Espacios hasta el final de la página
    for _ in range(15):
        story.append(Spacer(1, 12))
    
    # Línea horizontal al final
    story.append(Spacer(1, 12))
    
    # Salto de página
    story.append(PageBreak())
    
    # PÁGINA 2 - CONTENIDO PRINCIPAL
    # Título repetido
    story.append(Paragraph(titulo_principal, title_style))
    story.append(Paragraph(acta.razon_social.upper(), title_style))
    story.append(Spacer(1, 12))
    
    # Párrafo de apertura
    hora_inicio = acta.hora_inicio.strftime('%H:%M') if acta.hora_inicio else "[HORA]"
    parrafo_apertura = f"""Siendo las {hora_inicio} horas del {fecha_formateada}, se reunieron en el domicilio social de {acta.razon_social} (en adelante referida indistintamente por su denominación social, "<u>la Sociedad</u>") ubicado en {acta.lugar}, las personas cuyos nombres aparecen en la lista de asistencia (en adelante y en conjunto denominados los "Asambleístas"), que firmada por el Secretario, Escrutador y los accionistas, se agrega a la presente acta como "Anexo Uno", para celebrar una Asamblea General {acta.get_tipo_asamblea_display() if hasattr(acta, 'get_tipo_asamblea_display') else acta.tipo_asamblea} de Accionistas de la Sociedad (en adelante la "Asamblea"), para la cual {'no se requirió convocatoria alguna' if acta.convocatoria_omitida else 'se publicó la convocatoria correspondiente'}, al encontrarse presente el {acta.porcentaje_capital_presente}% de las acciones con derecho a voto representativas del capital social de la Sociedad de conformidad con lo dispuesto en los estatutos sociales."""
    
    story.append(Paragraph(parrafo_apertura, justified_style))
    
    # Párrafo de funcionarios
    parrafo_funcionarios = f"""De conformidad con los estatutos sociales, fungió como Presidente de la Asamblea {acta.presidente} y como Secretario actuó {acta.secretario}."""
    story.append(Paragraph(parrafo_funcionarios, justified_style))
    
    # Comisario si existe
    if acta.comisario:
        parrafo_comisario = f"Se hace constar la presencia del Comisario de la Sociedad, {acta.comisario}."
        story.append(Paragraph(parrafo_comisario, justified_style))
    
    # Escrutador
    if acta.escrutador:
        parrafo_escrutador = f"""El Presidente de la Asamblea designó como Escrutador a {acta.escrutador}, quien después de aceptar su cargo, procedió a realizar el escrutinio, para lo cual revisó el Libro de Registro de Acciones de la Sociedad y preparó la Lista de Asistencia, misma que aquí se tiene por íntegramente reproducida como si a la letra se insertase y en la cual, se hace constar que estuvieron presentes y representadas en la Asamblea la totalidad de las acciones con derecho a voto, representativas del capital social de la Sociedad."""
        story.append(Paragraph(parrafo_escrutador, justified_style))
    
    # Instalación de la asamblea
    fundamento = acta.fundamento_convocatoria if acta.fundamento_convocatoria else "los estatutos sociales y la Ley General de Sociedades Mercantiles"
    parrafo_instalacion = f"""Con base en el escrutinio debidamente certificado por el Secretario y Escrutador, el Presidente declaró formalmente instalada la Asamblea, {'no obstante que no se publicó la convocatoria respectiva' if acta.convocatoria_omitida else 'habiendo cumplido con los requisitos de convocatoria'}, ello en virtud de encontrarse presente el {acta.porcentaje_capital_presente}% de las acciones en circulación con derecho a voto, al tenor de lo dispuesto por {fundamento}."""
    story.append(Paragraph(parrafo_instalacion, justified_style))
    
    # Método de votación
    metodo_votacion = acta.get_metodo_votacion_general_display() if hasattr(acta, 'get_metodo_votacion_general_display') else acta.metodo_votacion_general
    parrafo_votacion = f"A continuación, el Presidente preguntó a los Asambleístas si estaban de acuerdo en que la votación se hiciese en forma {metodo_votacion.lower()}, a lo cual estos contestaron unánimemente en sentido afirmativo."
    story.append(Paragraph(parrafo_votacion, justified_style))
    
    # Orden del día
    parrafo_orden = "Acto seguido y a solicitud del Presidente de la Asamblea, el Secretario dio lectura al \"Orden del Día\", mismo que sometió a la aprobación de los Asambleístas y que a continuación se transcribe:"
    story.append(Paragraph(parrafo_orden, justified_style))
    
    story.append(Spacer(1, 12))
    story.append(Paragraph("ORDEN DEL DÍA", subtitle_style))
    story.append(Spacer(1, 12))
    
    # Incluir orden del día desde JSON
    if hasattr(acta, 'orden_dia_json') and acta.orden_dia_json:
        try:
            orden_dia_data = acta.orden_dia_json
            if isinstance(orden_dia_data, list):
                numeros_romanos = ['I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX', 'X']
                for i, punto in enumerate(orden_dia_data):
                    if i < len(numeros_romanos):
                        numero = numeros_romanos[i]
                        texto_punto = f"{numero}.&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;{punto}"
                        story.append(Paragraph(texto_punto, order_style))
                        story.append(Spacer(1, 6))
        except:
            pass
        story.append(Spacer(1, 8))
    story.append(Spacer(1, 15))
    
    # Párrafo de cierre del orden del día
    parrafo_cierre_orden = """Los Asambleístas después de escuchar el Orden del Día antes transcrito y de haberlo aprobado por unanimidad, procedieron a discutir ampliamente todos y cada uno de los asuntos contenidos en el mismo."""
    story.append(Paragraph(parrafo_cierre_orden, justified_style))
    story.append(Spacer(1, 20))
    
    # DESARROLLO DE LOS PUNTOS DEL ORDEN DEL DÍA
    story.append(Paragraph("DESARROLLO DE LOS PUNTOS DEL ORDEN DEL DÍA", subtitle_style))
    story.append(Spacer(1, 15))
    
    # Incluir resoluciones si están disponibles
    if hasattr(acta, 'resoluciones_json') and acta.resoluciones_json:
        try:
            resoluciones_data = acta.resoluciones_json
            if isinstance(resoluciones_data, list):
                resoluciones_texto = '\n\n'.join(resoluciones_data)
            elif isinstance(resoluciones_data, dict):
                resoluciones_texto = '\n\n'.join([f"{k}: {v}" for k, v in resoluciones_data.items()])
            else:
                resoluciones_texto = str(resoluciones_data)
        except:
            resoluciones_texto = str(acta.resoluciones_json)
        
        story.append(Paragraph("RESOLUCIONES ADOPTADAS", subtitle_style))
        story.append(Paragraph(resoluciones_texto, justified_style))
        story.append(Spacer(1, 20))
    
    # Incluir nombramientos si están disponibles
    if hasattr(acta, 'nombramientos_json') and acta.nombramientos_json:
        try:
            nombramientos_data = acta.nombramientos_json
            if isinstance(nombramientos_data, list):
                nombramientos_texto = '\n\n'.join(nombramientos_data)
            elif isinstance(nombramientos_data, dict):
                nombramientos_texto = '\n\n'.join([f"{k}: {v}" for k, v in nombramientos_data.items()])
            else:
                nombramientos_texto = str(nombramientos_data)
        except:
            nombramientos_texto = str(acta.nombramientos_json)
        
        story.append(Paragraph("NOMBRAMIENTOS Y DESIGNACIONES", subtitle_style))
        story.append(Paragraph(nombramientos_texto, justified_style))
        story.append(Spacer(1, 20))
    
    # Incluir información de dividendos si está disponible
    if hasattr(acta, 'dividendos_json') and acta.dividendos_json:
        try:
            dividendos_data = acta.dividendos_json
            if isinstance(dividendos_data, list):
                dividendos_texto = '\n\n'.join(dividendos_data)
            elif isinstance(dividendos_data, dict):
                dividendos_texto = '\n\n'.join([f"{k}: {v}" for k, v in dividendos_data.items()])
            else:
                dividendos_texto = str(dividendos_data)
        except:
            dividendos_texto = str(acta.dividendos_json)
        
        story.append(Paragraph("DECRETO DE DIVIDENDOS", subtitle_style))
        story.append(Paragraph(dividendos_texto, justified_style))
        story.append(Spacer(1, 20))
    
    # Incluir delegados si están disponibles
    if hasattr(acta, 'delegados_json') and acta.delegados_json:
        try:
            delegados_data = acta.delegados_json
            if isinstance(delegados_data, list):
                delegados_texto = '\n\n'.join(delegados_data)
            elif isinstance(delegados_data, dict):
                delegados_texto = '\n\n'.join([f"{k}: {v}" for k, v in delegados_data.items()])
            else:
                delegados_texto = str(delegados_data)
        except:
            delegados_texto = str(acta.delegados_json)
        
        story.append(Paragraph("DELEGADOS ESPECIALES", subtitle_style))
        story.append(Paragraph(delegados_texto, normal_style))
        story.append(Spacer(1, 20))
    
    # CLAUSURA DE LA ASAMBLEA
    story.append(Paragraph("CLAUSURA DE LA ASAMBLEA", subtitle_style))
    hora_cierre = acta.hora_cierre.strftime('%H:%M') if hasattr(acta, 'hora_cierre') and acta.hora_cierre else "[HORA DE CIERRE]"
    parrafo_clausura = f"""No habiendo más asuntos que tratar, el Presidente de la Asamblea declaró clausurada la sesión siendo las {hora_cierre} horas del mismo día, procediendo los asistentes a firmar la presente acta en señal de conformidad."""
    story.append(Paragraph(parrafo_clausura, normal_style))
    story.append(Spacer(1, 30))
    
    # FIRMAS
    story.append(Paragraph("FIRMAS", subtitle_style))
    story.append(Spacer(1, 20))
    
    # Crear tabla de firmas
    firmas_data = []
    
    # Presidente
    presidente_nombre = acta.presidente if hasattr(acta, 'presidente') and acta.presidente else "[PRESIDENTE]"
    firmas_data.append(["_" * 40, "_" * 40])
    firmas_data.append([presidente_nombre, "PRESIDENTE DE LA ASAMBLEA"])
    firmas_data.append(["", ""])
    
    # Secretario
    secretario_nombre = acta.secretario if hasattr(acta, 'secretario') and acta.secretario else "[SECRETARIO]"
    firmas_data.append(["_" * 40, "_" * 40])
    firmas_data.append([secretario_nombre, "SECRETARIO DE LA ASAMBLEA"])
    firmas_data.append(["", ""])
    
    # Escrutador si existe
    if hasattr(acta, 'escrutador') and acta.escrutador:
        firmas_data.append(["_" * 40, "_" * 40])
        firmas_data.append([acta.escrutador, "ESCRUTADOR"])
        firmas_data.append(["", ""])
    
    # Comisario si existe
    if hasattr(acta, 'comisario') and acta.comisario:
        firmas_data.append(["_" * 40, "_" * 40])
        firmas_data.append([acta.comisario, "COMISARIO"])
    
    # Crear tabla de firmas
    firmas_table = Table(firmas_data, colWidths=[8*cm, 8*cm])
    firmas_table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('TOPPADDING', (0, 0), (-1, -1), 12),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
    ]))
    
    story.append(firmas_table)
    story.append(Spacer(1, 30))
    
    # ANEXOS si están disponibles
    if hasattr(acta, 'anexos_json') and acta.anexos_json:
        story.append(PageBreak())
        story.append(Paragraph("ANEXOS", subtitle_style))
        story.append(Spacer(1, 15))
        
        try:
            anexos_data = acta.anexos_json
            if isinstance(anexos_data, list):
                anexos_texto = '\n\n'.join(anexos_data)
            elif isinstance(anexos_data, dict):
                anexos_texto = '\n\n'.join([f"{k}: {v}" for k, v in anexos_data.items()])
            else:
                anexos_texto = str(anexos_data)
        except:
            anexos_texto = str(acta.anexos_json)
        
        story.append(Paragraph(anexos_texto, normal_style))
        story.append(Spacer(1, 20))
    
    # LISTA DE ASISTENTES si está disponible
    if hasattr(acta, 'asistentes_json') and acta.asistentes_json:
        story.append(PageBreak())
        story.append(Paragraph("ANEXO I - LISTA DE ASISTENCIA", subtitle_style))
        story.append(Spacer(1, 15))
        
        try:
            asistentes_data = acta.asistentes_json
            if isinstance(asistentes_data, list):
                asistentes_texto = '\n\n'.join(asistentes_data)
            elif isinstance(asistentes_data, dict):
                asistentes_texto = '\n\n'.join([f"{k}: {v}" for k, v in asistentes_data.items()])
            else:
                asistentes_texto = str(asistentes_data)
        except:
            asistentes_texto = str(acta.asistentes_json)
        
        story.append(Paragraph(asistentes_texto, normal_style))
        story.append(Spacer(1, 20))
    
    # Información adicional si está disponible
    if hasattr(acta, 'observaciones') and acta.observaciones:
        story.append(Spacer(1, 20))
        story.append(Paragraph("OBSERVACIONES ADICIONALES", subtitle_style))
        story.append(Paragraph(acta.observaciones, normal_style))
    
    story.append(Spacer(1, 40))
    
    # Pie de documento
    story.append(Paragraph("Documento generado por el Sistema de Gestión de Documentos Legales OLEA", center_style))
    fecha_generacion = f"Generado el {fecha_completa}"
    story.append(Paragraph(fecha_generacion, center_style))
    
    # Construir PDF
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()
def descargar_pdf_asamblea(request, pk):
    """Vista para descargar PDF de Acta de Asamblea"""
    acta = get_object_or_404(ActaAsamblea, pk=pk, usuario=request.user)
    try:
        pdf = generar_pdf_acta_asamblea(acta)
        response = HttpResponse(pdf, content_type='application/pdf')
        filename = f'acta_asamblea_{acta.pk}_{acta.fecha.strftime("%Y%m%d") if acta.fecha else "sin_fecha"}.pdf'
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        messages.success(request, f'PDF generado: {filename}')
        return response
    except Exception as e:
        messages.error(request, f'Error al generar PDF: {e}')
        return redirect('documentos:detalle', pk=pk)


def generar_docx_pagare(pagare):
    """Genera un documento DOCX basado en template con reemplazos de datos del pagaré"""
    print(f"DEBUG: Iniciando generación DOCX para pagaré ID: {pagare.pk}")
    
    if not DOCX_AVAILABLE:
        raise ImportError("La librería python-docx no está disponible")
    
    # Ruta del template
    template_path = os.path.join(
        settings.BASE_DIR.parent, 
        "DOCUMENTOS OLEA ABOGADOS", 
        "Pagarés", 
        "PAGARE_PLACE.docx"
    )
    
    print(f"DEBUG: Buscando template en: {template_path}")
    print(f"DEBUG: Template existe: {os.path.exists(template_path)}")
    
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template no encontrado en: {template_path}")
    
    # Cargar el template
    print("DEBUG: Cargando template...")
    doc = Document(template_path)
    print(f"DEBUG: Template cargado. Párrafos encontrados: {len(doc.paragraphs)}")
    
    # Preparar datos para reemplazo
    fecha_emision_formateada = pagare.fecha_emision.strftime("%d de %B de %Y") if pagare.fecha_emision else "[FECHA_EMISION]"
    
    # Convertir mes a español
    meses = {
        'January': 'enero', 'February': 'febrero', 'March': 'marzo',
        'April': 'abril', 'May': 'mayo', 'June': 'junio',
        'July': 'julio', 'August': 'agosto', 'September': 'septiembre',
        'October': 'octubre', 'November': 'noviembre', 'December': 'diciembre'
    }
    for eng, esp in meses.items():
        fecha_emision_formateada = fecha_emision_formateada.replace(eng, esp)
    
    # Diccionario de reemplazos
    replacements = {
        '{{lugar_emision}}': pagare.lugar_emision or '[LUGAR_EMISION]',
        '{{fecha_emision}}': fecha_emision_formateada,
        '{{nombre_acreedor}}': pagare.acreedor_nombre or '[NOMBRE_ACREEDOR]',
        '{{NOMBRE_ACREEDOR}}': (pagare.acreedor_nombre.upper() if pagare.acreedor_nombre else '[NOMBRE_ACREEDOR]'),
        '{{rfc_acreedor}}': pagare.acreedor_rfc or '[RFC_ACREEDOR]',
        '{{domicilio_acreedor}}': pagare.acreedor_domicilio or '[DOMICILIO_ACREEDOR]',
        '{{nombre_deudor}}': pagare.deudor_nombre or '[NOMBRE_DEUDOR]',
        '{{NOMBRE_DEUDOR}}': (pagare.deudor_nombre.upper() if pagare.deudor_nombre else '[NOMBRE_DEUDOR]'),
        '{{rfc_deudor}}': pagare.deudor_rfc or '[RFC_DEUDOR]',
        '{{representante_deudor}}': pagare.deudor_representante or '[REPRESENTANTE_DEUDOR]',
        '{{domicilio_deudor}}': pagare.deudor_domicilio or '[DOMICILIO_DEUDOR]',
        '{{monto_principal}}': f"${pagare.monto_numeric:,.2f}" if pagare.monto_numeric else '[MONTO_PRINCIPAL]',
        '{{moneda}}': pagare.moneda or '[MONEDA]',
        '{{concepto_pagare}}': pagare.concepto or '[CONCEPTO_PAGARE]',
        '{{monto_principal_texto}}': pagare.monto_literal or '[MONTO_PRINCIPAL_TEXTO]',
        '{{tipo_pago}}': pagare.get_tipo_pago_display() if pagare.tipo_pago else '[TIPO_PAGO]',
        '{{numero_pagos}}': str(pagare.num_pagos) if pagare.num_pagos else '[NUMERO_PAGOS]',
        '{{numero_pagos_letra}}': (num2words(pagare.num_pagos, lang="es").capitalize()if pagare.num_pagos is not None else ''),   
        '{{periodicidad}}': pagare.get_periodicidad_display() if pagare.periodicidad else '[PERIODICIDAD]',
        '{{lugar_pago}}': pagare.lugar_pago or '[LUGAR_PAGO]',
        '{{forma_pago}}': pagare.forma_pago or '[FORMA_PAGO]',
        '{{tasa_interes_ordinario}}': f"{pagare.tasa_interes_ordinario}%" if pagare.tasa_interes_ordinario else '[TASA_INTERES_ORDINARIO]',
        '{{tasa_interes_ordinario_letra}}': num2words(pagare.tasa_interes_ordinario, lang="es").capitalize(),
        '{{tasa_interes_moratorio}}': f"{pagare.tasa_interes_moratorio}%" if pagare.tasa_interes_moratorio else '[TASA_INTERES_MORATORIO]',
        '{{tasa_interes_moratorio_letra}}': num2words(pagare.tasa_interes_moratorio, lang="es").capitalize(),
        '{{base_calculo_dias}}': str(pagare.base_intereses) if pagare.base_intereses else '[BASE_CALCULO_DIAS]',
        '{{base_calculo_dias_letra}}': num2words(pagare.base_intereses, lang="es").capitalize(),
        '{{gastos_administracion}}': f"${pagare.gastos_admon:,.2f}" if pagare.gastos_admon else '[GASTOS_ADMINISTRACION]',
        '{{gastos_administracion_letra}}': num2words(pagare.gastos_admon, lang="es").capitalize(),
        '{{permitir_prepago}}': 'Sí' if pagare.prepago_permitido else 'No',
        '{{dias_aviso_prepago}}': str(pagare.dias_aviso_prepago) if pagare.dias_aviso_prepago else '[DIAS_AVISO_PREPAGO]',
        '{{dias_aviso_prepago_letra}}': num2words(pagare.dias_aviso_prepago, lang="es").capitalize(),
        '{{condicion_prepago}}': pagare.condicion_prepago or '[CONDICION_PREPAGO]',
        '{{tiene_garantia_aval}}': 'Sí' if pagare.tiene_garantia else 'No',
        '{{nombre_aval}}': pagare.aval_nombre or '[NOMBRE_AVAL]',
        '{{NOMBRE_AVAL}}': (pagare.aval_nombre.upper() if pagare.aval_nombre else '[NOMBRE_AVAL]'),
        '{{descripcion_garantia}}': pagare.descripcion_garantia or '[DESCRIPCION_GARANTIA]',
        '{{domicilio_aval}}': pagare.aval_domicilio or '[DOMICILIO_AVAL]',
        '{{jurisdiccion}}': pagare.jurisdiccion or '[JURISDICCION]',
        '{{ley_aplicable}}': pagare.ley_aplicable or '[LEY_APLICABLE]',
        '{{eventos_incumplimiento}}': pagare.eventos_incumplimiento or '[EVENTOS_INCUMPLIMIENTO]',
        '{{clausula_aceleracion}}': pagare.clausula_aceleracion or '[CLAUSULA_ACELERACION]',
    }
    
    # Función auxiliar para reemplazar texto completo en párrafos
    def replace_in_paragraph(paragraph, replacements):
        # Obtener todo el texto del párrafo
        full_text = paragraph.text
        
        # Debug: imprimir texto encontrado
        if full_text.strip():
            print(f"DEBUG: Párrafo encontrado: '{full_text}'")
        
        # Verificar si hay placeholders en el párrafo
        has_placeholder = False
        found_placeholders = []
        for placeholder in replacements.keys():
            if placeholder in full_text:
                has_placeholder = True
                found_placeholders.append(placeholder)
        
        if found_placeholders:
            print(f"DEBUG: Placeholders encontrados: {found_placeholders}")
        
        if not has_placeholder:
            return
        
        # Realizar reemplazos en el texto completo
        new_text = full_text
        for placeholder, value in replacements.items():
            if placeholder in new_text:
                print(f"DEBUG: Reemplazando '{placeholder}' con '{value}'")
                new_text = new_text.replace(placeholder, str(value))
        
        print(f"DEBUG: Texto original: '{full_text}'")
        print(f"DEBUG: Texto nuevo: '{new_text}'")
        
        # Si el texto cambió, reemplazar todo el párrafo
        if new_text != full_text:
            # Limpiar todos los runs existentes
            for run in paragraph.runs:
                run.text = ""
            
            # Agregar el nuevo texto al primer run
            if paragraph.runs:
                paragraph.runs[0].text = new_text
            else:
                paragraph.add_run(new_text)
            print(f"DEBUG: Párrafo actualizado exitosamente")
    
    # Reemplazar en párrafos
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, replacements)
    
    # Reemplazar en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, replacements)
    
    # Reemplazar en headers y footers
    for section in doc.sections:
        # Header
        header = section.header
        for paragraph in header.paragraphs:
            replace_in_paragraph(paragraph, replacements)
        
        # Footer
        footer = section.footer
        for paragraph in footer.paragraphs:
            replace_in_paragraph(paragraph, replacements)
    
    # Guardar en BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generar_docx_acta_asamblea(acta):
    """Genera un documento DOCX basado en template con reemplazos de datos del acta de asamblea"""
    if not DOCX_AVAILABLE:
        raise ImportError("La librería python-docx no está disponible")
    
    # Ruta del template
    template_path = os.path.join(
        settings.BASE_DIR.parent, 
        "DOCUMENTOS OLEA ABOGADOS", 
        "Actas de Asambleas", 
        "Acta de Asamblea PLACE.docx"
    )
    
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template no encontrado en: {template_path}")
    
    # Cargar el template
    doc = Document(template_path)
    
    # Preparar datos para reemplazo
    fecha_formateada = acta.fecha.strftime("%d de %B de %Y") if acta.fecha else "[FECHA]"
    # Convertir mes a español
    meses = {
        'January': 'enero', 'February': 'febrero', 'March': 'marzo',
        'April': 'abril', 'May': 'mayo', 'June': 'junio',
        'July': 'julio', 'August': 'agosto', 'September': 'septiembre',
        'October': 'octubre', 'November': 'noviembre', 'December': 'diciembre'
    }
    for eng, esp in meses.items():
        fecha_formateada = fecha_formateada.replace(eng, esp)
    
    # Crear diccionario de reemplazos con manejo de campos vacíos
    replacements = {
        "{{razon_social}}": acta.razon_social or "[RAZÓN SOCIAL]",
        "{{RAZON_SOCIAL}}": acta.razon_social.upper() if acta.razon_social else "[RAZÓN SOCIAL]",
        "{{ABREVIATURA_SOCIEDAD}}": "".join([p[0] for p in acta.razon_social.strip().split() if p.upper() not in {"S", "S.A.", "S.A", "DE", "C.V.", "C.V", "RL", "R.L.", "R.L"}]).upper() if acta.razon_social else "CAS",
        "{{tipo_asamblea}}": acta.get_tipo_asamblea_display() if hasattr(acta, 'get_tipo_asamblea_display') else (acta.tipo_asamblea or "[TIPO DE ASAMBLEA]"),
        "{{TIPO_ASAMBLEA}}": acta.get_tipo_asamblea_display().upper() if hasattr(acta, 'get_tipo_asamblea_display') and acta.tipo_asamblea else (acta.tipo_asamblea.upper() if acta.tipo_asamblea else "[TIPO DE ASAMBLEA]"),
        "{{caracter}}": acta.caracter or "[CARÁCTER]",
        "{{CARACTER}}": acta.caracter.upper() if acta.caracter else "[CARÁCTER]",
        "{{fecha}}": fecha_formateada,
        "{{hora_inicio}}": acta.hora_inicio.strftime("%H:%M") if acta.hora_inicio else "[HORA DE INICIO]",
        "{{hora_cierre}}": acta.hora_cierre.strftime("%H:%M") if acta.hora_cierre else "[HORA DE CIERRE]",
        "{{lugar}}": acta.lugar or "[LUGAR]",
        "{{porcentage_capital_presente}}": f"{acta.porcentaje_capital_presente}%" if acta.porcentaje_capital_presente else "[PORCENTAJE CAPITAL PRESENTE]",
        "{{presidente}}": acta.presidente or "[PRESIDENTE]",
        "{{secretario}}": acta.secretario or "[SECRETARIO]",
        "{{escrutador}}": acta.escrutador or "[ESCRUTADOR]",
        "{{comisario}}": acta.comisario or "[COMISARIO]"
    }
    
    # Función auxiliar para reemplazar texto completo en párrafos
    def replace_in_paragraph(paragraph, replacements):
        # Obtener todo el texto del párrafo
        full_text = paragraph.text
        
        # Verificar si hay placeholders en el párrafo
        has_placeholder = False
        for placeholder in replacements.keys():
            if placeholder in full_text:
                has_placeholder = True
                break
        
        if not has_placeholder:
            return
        
        # Realizar reemplazos en el texto completo
        new_text = full_text
        for placeholder, value in replacements.items():
            if placeholder in new_text:
                new_text = new_text.replace(placeholder, str(value))
        
        # Si el texto cambió, reemplazar todo el párrafo
        if new_text != full_text:
            # Limpiar todos los runs existentes
            for run in paragraph.runs:
                run.text = ""
            
            # Agregar el nuevo texto al primer run
            if paragraph.runs:
                paragraph.runs[0].text = new_text
            else:
                paragraph.add_run(new_text)
    
    # Procesar ORDENES_Y_RESOLUCIONES si existe el placeholder
    try:
        from .docx_blocks import build_ordenes_con_resoluciones, inject_ordenes_y_resoluciones
        
        # Verificar si existe el placeholder en el documento
        placeholder_found = False
        for paragraph in doc.paragraphs:
            if "{{ORDENES_Y_RESOLUCIONES}}" in paragraph.text:
                placeholder_found = True
                break
        
        if not placeholder_found:
            # Verificar en tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if "{{ORDENES_Y_RESOLUCIONES}}" in paragraph.text:
                                placeholder_found = True
                                break
                        if placeholder_found:
                            break
                    if placeholder_found:
                        break
        
        if placeholder_found:
            # Construir las órdenes con resoluciones
            ordenes = build_ordenes_con_resoluciones(
                acta.orden_dia_json, 
                acta.resoluciones_json
            )
            
            # Texto constante para Asambleas (similar al de Consejo pero adaptado)
            texto_constante = (
                "Los Asambleístas después de escuchar el Orden del Día antes trascrito y de haberlo aprobado por unanimidad, procedieron a discutir ampliamente todos y cada uno de los asuntos contenidos en el mismo, como sigue:"
            )
            
            # Inyectar el contenido formateado
            inject_ordenes_y_resoluciones(
                doc, 
                "{{ORDENES_Y_RESOLUCIONES}}", 
                ordenes, 
                texto_constante
            )
            
            print(f"DEBUG: ORDENES_Y_RESOLUCIONES procesado exitosamente para Acta de Asamblea. {len(ordenes)} órdenes encontradas.")
        else:
            print("DEBUG: Placeholder {{ORDENES_Y_RESOLUCIONES}} no encontrado en template de Asamblea.")
            
    except Exception as e:
        print(f"ERROR procesando ORDENES_Y_RESOLUCIONES en Asamblea: {e}")
        # Continuar con el procesamiento normal si hay error
    
    # Aplicar reemplazos en documento principal
    replace_in_xml_textnodes(doc.part, replacements)
    
    # Reemplazar en párrafos
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, replacements)
    
    # Reemplazar en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, replacements)
    
    # Aplicar reemplazos en headers y footers
    for section in doc.sections:
        for header in (section.header, section.footer):
            replace_in_xml_textnodes(header.part, replacements)
    
    # Guardar en memoria
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()


@login_required
def descargar_docx_acta_asamblea(request, pk):
    """Vista para descargar DOCX de Acta de Asamblea"""
    acta = get_object_or_404(ActaAsamblea, pk=pk, usuario=request.user)
    try:
        docx_content = generar_docx_acta_asamblea(acta)
        response = HttpResponse(
            docx_content, 
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        filename = f'acta_asamblea_{acta.pk}_{acta.fecha.strftime("%Y%m%d") if acta.fecha else "sin_fecha"}.docx'
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response
    except Exception as e:
        messages.error(request, f'Error al generar DOCX: {e}')
        return redirect('documentos:detalle', pk=pk)


@login_required
def descargar_docx_pagare(request, pk):
    """Vista para descargar DOCX de Pagaré"""
    pagare = get_object_or_404(Pagare, pk=pk, usuario=request.user)
    
    # Debug: verificar datos del pagaré
    print(f"DEBUG: Pagaré ID: {pagare.pk}")
    print(f"DEBUG: Lugar emisión: '{pagare.lugar_emision}'")
    print(f"DEBUG: Fecha emisión: '{pagare.fecha_emision}'")
    print(f"DEBUG: Acreedor: '{pagare.acreedor_nombre}'")
    print(f"DEBUG: Deudor: '{pagare.deudor_nombre}'")
    print(f"DEBUG: Monto: '{pagare.monto_numeric}'")
    
    try:
        docx_content = generar_docx_pagare(pagare)
        response = HttpResponse(
            docx_content, 
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        filename = f'pagare_{pagare.pk}_{pagare.fecha_emision.strftime("%Y%m%d")}.docx'
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        print(f"DEBUG: DOCX generado exitosamente: {filename}")
        return response
    except Exception as e:
        print(f"DEBUG: Error al generar DOCX: {e}")
        import traceback
        traceback.print_exc()
        messages.error(request, f'Error al generar DOCX: {e}')
        return redirect('documentos:detalle_pagare', pk=pk)


@login_required
def descargar_docx_prenda(request, pk):
    """Vista para descargar DOCX de Contrato de Prenda sobre Acciones"""
    contrato = get_object_or_404(ContratoPrendaAcciones, pk=pk, usuario=request.user)
    try:
        return generar_docx_prenda(contrato)
    except Exception as e:
        messages.error(request, f'Error al generar DOCX: {e}')
        return redirect('documentos:detalle_contrato_prenda', pk=pk)
