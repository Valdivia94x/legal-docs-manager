from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from django.conf import settings
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
import os
import re
from .models import ContratoCredito
from num2words import num2words
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def formatear_fecha(fecha):
    """Convierte una fecha a formato español legible"""
    if not fecha:
        return "[FECHA]"
    fecha_str = fecha.strftime("%d de %B de %Y")
    meses = {
        'January': 'enero', 'February': 'febrero', 'March': 'marzo',
        'April': 'abril', 'May': 'mayo', 'June': 'junio',
        'July': 'julio', 'August': 'agosto', 'September': 'septiembre',
        'October': 'octubre', 'November': 'noviembre', 'December': 'diciembre'
    }
    for eng, esp in meses.items():
        fecha_str = fecha_str.replace(eng, esp)
    return fecha_str

def numero_a_texto(numero):
    """Convierte un número a texto en español"""
    if numero is None:
        return ""
    try:
        return num2words(float(numero), lang='es')
    except:
        return str(numero)

def fecha_a_texto(fecha):
    meses = [
        "enero", "febrero", "marzo", "abril", "mayo", "junio",
        "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"
    ]
    dia = num2words(fecha.day, lang="es").capitalize()
    mes = meses[fecha.month - 1]
    anio = num2words(fecha.year, lang="es")
    return f"{dia} de {mes} de {anio}"

def descargar_docx_contrato_credito(request, pk):
    """
    Genera y descarga un archivo DOCX de Contrato de Crédito con placeholders reemplazados
    """
    # Obtener el objeto ContratoCredito
    contrato = get_object_or_404(ContratoCredito, pk=pk, usuario=request.user)
    
    # Ruta al template DOCX
    template_path = os.path.join(
        settings.BASE_DIR.parent, 
        "DOCUMENTOS OLEA ABOGADOS", 
        "Contratos de Crédito", 
        "Contrato de Credito PLACE.docx"
    )
    
    if not os.path.exists(template_path):
        return HttpResponse("Template file not found", status=404)
    
    try:
        # Cargar el documento template
        doc = Document(template_path)
        
        # Diccionario de reemplazos - mapeando campos del modelo a placeholders
        replacements = {
            # Datos Generales del Contrato
            '{{fecha_contrato}}': formatear_fecha(contrato.fecha_contrato),
            '{{fecha_contrato_texto}}': fecha_a_texto(contrato.fecha_contrato),
            '{{lugar_contrato}}': contrato.lugar_contrato or '',
            
            # Parte Acreditante
            '{{acreditante_razon_social}}': contrato.acreditante_razon_social or '',
            '{{ACREDITANTE_RAZON_SOCIAL}}': contrato.acreditante_razon_social.upper() if contrato.acreditante_razon_social else '',
            '{{acreditante_forma_legal}}': contrato.acreditante_forma_legal or '',
            '{{acreditante_representante}}': contrato.acreditante_representante or '',
            '{{ACREDITANTE_REPRESENTANTE}}': contrato.acreditante_representante.upper() if contrato.acreditante_representante else '',
            
            # Constitución original del acreditante
            '{{acreditante_deed_constitucion_numero}}': contrato.acreditante_deed_constitucion_numero or '',
            '{{acreditante_deed_constitucion_fecha}}': formatear_fecha(contrato.acreditante_deed_constitucion_fecha),
            '{{acreditante_notario_constitucion}}': contrato.acreditante_notario_constitucion or '',
            '{{acreditante_registro_constitucion_folio}}': contrato.acreditante_registro_constitucion_folio or '',
            
            # Cambio a S.A.P.I. del acreditante
            '{{acreditante_deed_prom_inv_numero}}': contrato.acreditante_deed_prom_inv_numero or '',
            '{{acreditante_deed_prom_inv_fecha}}': formatear_fecha(contrato.acreditante_deed_prom_inv_fecha),
            '{{acreditante_notario_prom_inv}}': contrato.acreditante_notario_prom_inv or '',
            '{{acreditante_registro_prom_inv_folio}}': contrato.acreditante_registro_prom_inv_folio or '',
            
            # Poder del representante del acreditante
            '{{acreditante_deed_poder_numero}}': contrato.acreditante_deed_poder_numero or '',
            '{{acreditante_deed_poder_fecha}}': formatear_fecha(contrato.acreditante_deed_poder_fecha),
            '{{acreditante_notario_poder}}': contrato.acreditante_notario_poder or '',
            
            # Parte Acreditado
            '{{acreditado_razon_social_original}}': contrato.acreditado_razon_social_original or '',
            '{{ACREDITADO_RAZON_SOCIAL_ORIGINAL}}': contrato.acreditado_razon_social_original.upper() if contrato.acreditado_razon_social_original else '',
            
            # Constitución original del acreditado
            '{{acreditado_deed_constitucion_original_numero}}': contrato.acreditado_deed_constitucion_original_numero or '',
            '{{acreditado_deed_constitucion_original_fecha}}': formatear_fecha(contrato.acreditado_deed_constitucion_original_fecha),
            '{{acreditado_notario_constitucion_original}}': contrato.acreditado_notario_constitucion_original or '',
            '{{acreditado_registro_constitucion_original_folio}}': contrato.acreditado_registro_constitucion_original_folio or '',
            
            # Cambio de denominación social del acreditado
            '{{acreditado_deed_denominacion_cambio_numero}}': contrato.acreditado_deed_denominacion_cambio_numero or '',
            '{{acreditado_deed_denominacion_cambio_fecha}}': formatear_fecha(contrato.acreditado_deed_denominacion_cambio_fecha),
            '{{acreditado_notario_denominacion_cambio}}': contrato.acreditado_notario_denominacion_cambio or '',
            '{{acreditado_registro_denominacion_cambio_folio}}': contrato.acreditado_registro_denominacion_cambio_folio or '',
            
            # Poder del representante del acreditado
            '{{acreditado_deed_poder_numero}}': contrato.acreditado_deed_poder_numero or '',
            '{{acreditado_deed_poder_fecha}}': formatear_fecha(contrato.acreditado_deed_poder_fecha),
            '{{acreditado_notario_poder}}': contrato.acreditado_notario_poder or '',
            
            # Resoluciones del consejo del acreditado
            '{{acreditado_resoluciones_unani_fecha}}': formatear_fecha(contrato.acreditado_resoluciones_unani_fecha),
            
            # Términos del Crédito
            '{{monto_credito}}': f"${contrato.monto_credito:,.2f}" if contrato.monto_credito else '',
            '{{monto_credito_texto}}': contrato.monto_credito_texto or '',
            '{{MONTO_CREDITO_TEXTO}}': contrato.monto_credito_texto.upper() if contrato.monto_credito_texto else '',
            '{{numero_disposiciones}}': str(contrato.numero_disposiciones) if contrato.numero_disposiciones else '',
            
            # Disposiciones del crédito
            '{{disposicion1_fecha}}': formatear_fecha(contrato.disposicion1_fecha),
            '{{disposicion1_importe}}': f"${contrato.disposicion1_importe:,.2f}" if contrato.disposicion1_importe else '',
            '{{disposicion1_importe_texto}}': num2words(contrato.disposicion1_importe, lang='es') if contrato.disposicion1_importe else '',
            '{{disposicion2_fecha}}': formatear_fecha(contrato.disposicion2_fecha),
            '{{disposicion2_fecha_texto}}': fecha_a_texto(contrato.disposicion2_fecha),
            '{{disposicion2_importe}}': f"${contrato.disposicion2_importe:,.2f}" if contrato.disposicion2_importe else '',
            '{{disposicion3_fecha}}': formatear_fecha(contrato.disposicion3_fecha),
            '{{disposicion3_fecha_texto}}': fecha_a_texto(contrato.disposicion3_fecha),
            '{{disposicion3_importe}}': f"${contrato.disposicion3_importe:,.2f}" if contrato.disposicion3_importe else '',
            
            # Fechas de vencimiento y pagos
            '{{plazo_credito_fecha_vencimiento}}': formatear_fecha(contrato.plazo_credito_fecha_vencimiento),
            '{{pagos_intereses_fecha1}}': formatear_fecha(contrato.pagos_intereses_fecha1),
            '{{pagos_intereses_fecha2}}': formatear_fecha(contrato.pagos_intereses_fecha2),
            '{{pagos_intereses_fecha3}}': formatear_fecha(contrato.pagos_intereses_fecha3),
            '{{pago_principal_fecha}}': formatear_fecha(contrato.pago_principal_fecha),
            
            # Tasas de interés
            '{{tasa_interes_ordinaria}}': f"{contrato.tasa_interes_ordinaria}%" if contrato.tasa_interes_ordinaria else '',
            '{{tasa_interes_moratoria}}': f"{contrato.tasa_interes_moratoria}%" if contrato.tasa_interes_moratoria else '',
            
            # Información bancaria
            '{{banco_cuenta_numero}}': contrato.banco_cuenta_numero or '',
            '{{banco_nombre}}': contrato.banco_nombre or '',
            '{{banco_titular}}': contrato.banco_titular or '',
            '{{banco_clabe}}': contrato.banco_clabe or '',
            
            # Domicilios
            '{{domicilio_acreditante}}': contrato.domicilio_acreditante or '',
            '{{domicilio_acreditado}}': contrato.domicilio_acreditado or '',
            
            # Jurisdicción y ley aplicable
            '{{jurisdiccion}}': contrato.jurisdiccion or '',
            '{{ley_aplicable}}': contrato.ley_aplicable or '',
            
            # Información del Aval
            '{{aval_nombre}}': contrato.aval_nombre or '',
            '{{AVAL_NOMBRE}}': contrato.aval_nombre.upper() if contrato.aval_nombre else '',
            '{{aval_domicilio}}': contrato.aval_domicilio or '',
        }
        
        # Función auxiliar para reemplazar texto completo en párrafos
        def replace_in_paragraph(paragraph, replacements):
            # Obtener todo el texto del párrafo
            full_text = paragraph.text
            
            # Verificar si hay placeholders en el párrafo
            has_placeholder = False
            for placeholder in replacements.keys():
                if placeholder in full_text:
                    has_placeholder = True
                    break
            
            if not has_placeholder:
                return
            
            # Realizar reemplazos normales
            new_text = full_text
            for placeholder, value in replacements.items():
                if placeholder in new_text:
                    new_text = new_text.replace(placeholder, str(value))
            
            # Si el texto cambió, reemplazar todo el párrafo
            if new_text != full_text:
                # Limpiar todos los runs existentes
                for run in paragraph.runs:
                    run.text = ""
                
                # Agregar el nuevo texto al primer run
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                else:
                    paragraph.add_run(new_text)
        
        # Reemplazar placeholders en párrafos
        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph, replacements)
        
        # Reemplazar placeholders en tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph, replacements)
        
        # Reemplazar placeholders en headers y footers
        for section in doc.sections:
            # Header
            if section.header:
                for paragraph in section.header.paragraphs:
                    replace_in_paragraph(paragraph, replacements)
            
            # Footer
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    replace_in_paragraph(paragraph, replacements)
        
        # Preparar la respuesta HTTP
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        # Nombre del archivo de descarga
        filename = f"Contrato_Credito_{contrato.acreditado_razon_social_original.replace(' ', '_') if contrato.acreditado_razon_social_original else 'documento'}_{contrato.pk}.docx"
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        # Guardar el documento en la respuesta
        doc.save(response)
        
        return response
        
    except Exception as e:
        return HttpResponse(f"Error generating document: {str(e)}", status=500)
